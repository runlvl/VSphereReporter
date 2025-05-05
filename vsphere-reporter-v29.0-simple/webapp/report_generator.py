#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
VMware vSphere Reporter v29.0 - Report Generator

Dieses Modul stellt die ReportGenerator-Klasse zur Verfügung, 
die Berichte in verschiedenen Formaten generiert.
"""

import os
import sys
import logging
import json
import re
from datetime import datetime
import tempfile
import shutil

# Versuche, die Berichtsmodule zu importieren
try:
    # Für PDF-Berichte
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    # Für Word-Berichte
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    
    REPORT_MODULES_AVAILABLE = True
except ImportError as e:
    REPORT_MODULES_AVAILABLE = False
    print(f"Warnung: Berichtsmodule konnten nicht importiert werden: {e}")
    print("Einige Exportfunktionen sind möglicherweise nicht verfügbar.")

from webapp.utils.error_handler import ReportGenerationError, handle_vsphere_errors

# Logger konfigurieren
logger = logging.getLogger(__name__)

class ReportGenerator:
    """
    Generator für Berichte in verschiedenen Formaten.
    """
    
    def __init__(self, output_dir=None, vcenter_server=None, username=None, demo_mode=False):
        """
        Initialisiere den Report-Generator.
        
        Args:
            output_dir: Verzeichnis für die Ausgabe der Berichte (optional)
            vcenter_server: Name des vCenter-Servers für die Berichte
            username: Benutzername für die Berichte
            demo_mode: Gibt an, ob der Demo-Modus aktiviert ist
        """
        self.output_dir = output_dir or tempfile.gettempdir()
        
        # Stelle sicher, dass das Ausgabeverzeichnis existiert
        os.makedirs(self.output_dir, exist_ok=True)
        
        self.vcenter_server = vcenter_server or "Unbekannter Server"
        self.username = username or "Unbekannter Benutzer"
        self.demo_mode = demo_mode
        
        # Prüfe, ob die Berichtsmodule verfügbar sind
        if not REPORT_MODULES_AVAILABLE:
            logger.warning("Berichtsmodule nicht verfügbar. Einige Exportfunktionen sind nicht nutzbar.")
        
        logger.info(f"Report-Generator initialisiert. Ausgabeverzeichnis: {self.output_dir}")
        logger.info(f"Server: {self.vcenter_server}, Benutzer: {self.username}, Demo-Mode: {self.demo_mode}")
    
    @handle_vsphere_errors
    def generate_html_report(self, filename, data):
        """
        Generiert einen HTML-Bericht mit den gegebenen Daten.
        
        Args:
            filename: Dateiname für den Bericht
            data: Dictionary mit den Berichtsdaten
                - vmware_tools: Liste von VMware Tools-Daten
                - snapshots: Liste von Snapshot-Daten
                - orphaned_vmdks: Liste von verwaisten VMDK-Daten
                
        Returns:
            str: Pfad zur generierten HTML-Datei
        """
        try:
            logger.info(f"Generiere HTML-Bericht: {filename}")
            
            # Stelle sicher, dass der Dateiname auf .html endet
            if not filename.lower().endswith('.html'):
                filename += '.html'
            
            # Vollständiger Pfad zur Ausgabedatei
            output_file = os.path.join(self.output_dir, filename)
            
            # Timestamp für den Bericht
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # HTML-Template für den Bericht
            html_template = f"""<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>VMware vSphere Bericht - {timestamp}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1, h2, h3 {{
            color: #00355e;
        }}
        h1 {{
            border-bottom: 2px solid #00355e;
            padding-bottom: 10px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 30px;
        }}
        th {{
            background-color: #00355e;
            color: white;
            text-align: left;
            padding: 8px;
        }}
        td {{
            padding: 8px;
            border-bottom: 1px solid #ddd;
        }}
        tr:nth-child(even) {{
            background-color: #f9f9f9;
        }}
        .warning {{
            background-color: #fff3cd;
        }}
        .danger {{
            background-color: #f8d7da;
        }}
        .info-box {{
            background-color: #f3f3f3;
            border-left: 4px solid #00355e;
            padding: 15px;
            margin-bottom: 20px;
        }}
        .footer {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            font-size: 0.9em;
            color: #666;
        }}
        .logo {{
            display: block;
            max-width: 200px;
            margin-bottom: 20px;
        }}
    </style>
</head>
<body>
    <img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAMgAAAAyCAYAAAAZUZThAAAACXBIWXMAAAsTAAALEwEAmpwYAAAF62lUWHRYTUw6Y29tLmFkb2JlLnhtcAAAAAAAPD94cGFja2V0IGJlZ2luPSLvu78iIGlkPSJXNU0wTXBDZWhpSHpyZVN6TlRjemtjOWQiPz4gPHg6eG1wbWV0YSB4bWxuczp4PSJhZG9iZTpuczptZXRhLyIgeDp4bXB0az0iQWRvYmUgWE1QIENvcmUgNi4wLWMwMDIgNzkuMTY0MzYwLCAyMDIwLzAyLzEzLTAxOjA3OjIyICAgICAgICAiPiA8cmRmOlJERiB4bWxuczpyZGY9Imh0dHA6Ly93d3cudzMub3JnLzE5OTkvMDIvMjItcmRmLXN5bnRheC1ucyMiPiA8cmRmOkRlc2NyaXB0aW9uIHJkZjphYm91dD0iIiB4bWxuczp4bXA9Imh0dHA6Ly9ucy5hZG9iZS5jb20veGFwLzEuMC8iIHhtbG5zOmRjPSJodHRwOi8vcHVybC5vcmcvZGMvZWxlbWVudHMvMS4xLyIgeG1sbnM6cGhvdG9zaG9wPSJodHRwOi8vbnMuYWRvYmUuY29tL3Bob3Rvc2hvcC8xLjAvIiB4bWxuczp4bXBNTT0iaHR0cDovL25zLmFkb2JlLmNvbS94YXAvMS4wL21tLyIgeG1sbnM6c3RFdnQ9Imh0dHA6Ly9ucy5hZG9iZS5jb20veGFwLzEuMC9zVHlwZS9SZXNvdXJjZUV2ZW50IyIgeG1wOkNyZWF0b3JUb29sPSJBZG9iZSBQaG90b3Nob3AgMjEuMSAoV2luZG93cykiIHhtcDpDcmVhdGVEYXRlPSIyMDI1LTA0LTE1VDEwOjIyOjA3KzAyOjAwIiB4bXA6TW9kaWZ5RGF0ZT0iMjAyNS0wNC0xNVQxMDoyNzoxMCswMjowMCIgeG1wOk1ldGFkYXRhRGF0ZT0iMjAyNS0wNC0xNVQxMDoyNzoxMCswMjowMCIgZGM6Zm9ybWF0PSJpbWFnZS9wbmciIHBob3Rvc2hvcDpDb2xvck1vZGU9IjMiIHBob3Rvc2hvcDpJQ0NQcm9maWxlPSJzUkdCIElFQzYxOTY2LTIuMSIgeG1wTU06SW5zdGFuY2VJRD0ieG1wLmlpZDozOWM5OWQ2MC00MzU0LTc2NDMtOTY0ZC0yYmFiZDE0OTZkZDMiIHhtcE1NOkRvY3VtZW50SUQ9ImFkb2JlOmRvY2lkOnBob3Rvc2hvcDpkYzBlMDgxMC03ZTEwLWIyNDItODU3Ni0wMjc5ZTdmNjkzODMiIHhtcE1NOk9yaWdpbmFsRG9jdW1lbnRJRD0ieG1wLmRpZDo5OWFiYWMzYy1lZDM2LWI5NDYtYTYyNS1lOTJiNWZjODk5YmIiPiA8eG1wTU06SGlzdG9yeT4gPHJkZjpTZXE+IDxyZGY6bGkgc3RFdnQ6YWN0aW9uPSJjcmVhdGVkIiBzdEV2dDppbnN0YW5jZUlEPSJ4bXAuaWlkOjk5YWJhYzNjLWVkMzYtYjk0Ni1hNjI1LWU5MmI1ZmM4OTliYiIgc3RFdnQ6d2hlbj0iMjAyNS0wNC0xNVQxMDoyMjowNyswMjowMCIgc3RFdnQ6c29mdHdhcmVBZ2VudD0iQWRvYmUgUGhvdG9zaG9wIDIxLjEgKFdpbmRvd3MpIi8+IDxyZGY6bGkgc3RFdnQ6YWN0aW9uPSJzYXZlZCIgc3RFdnQ6aW5zdGFuY2VJRD0ieG1wLmlpZDozOWM5OWQ2MC00MzU0LTc2NDMtOTY0ZC0yYmFiZDE0OTZkZDMiIHN0RXZ0OndoZW49IjIwMjUtMDQtMTVUMTA6Mjc6MTArMDI6MDAiIHN0RXZ0OnNvZnR3YXJlQWdlbnQ9IkFkb2JlIFBob3Rvc2hvcCAyMS4xIChXaW5kb3dzKSIgc3RFdnQ6Y2hhbmdlZD0iLyIvPiA8L3JkZjpTZXE+IDwveG1wTU06SGlzdG9yeT4gPC9yZGY6RGVzY3JpcHRpb24+IDwvcmRmOlJERj4gPC94OnhtcG1ldGE+IDw/eHBhY2tldCBlbmQ9InIiPz7UhvwXAAALMElEQVR4nO2ce4xdVRXGf3PnTl87nbbYkZZOMVBAKVg1Cl58IcRQBJUYDaIIClGJGjUxKolBjfyh/uEjxkejSIgPVEAwqFExCigUWqEULUjbqcx0OjDTmXbm0Znp42bu8Y+z1529zr3nnjvTPjR8yeTec/Zee629zt5r7bX2PvdCRkZGRkZGRkZGRkZGRkZGRkZGRkZGRkZGRkZGRkZGxgTgZC7eFIvP3gNcAMwClpvyw8DvgUHztwNcCzQDC035g8AYgU35TlN+Ziz3YuATQIspe9SUnQtsj7Vj5QXgA+b6euBvwOvA6aZsCPgE8FdTtwt4r8E0/1kF3G/qXhQrXwq8aH6fG5O3A4+Z339vbDMf+Dnw+ljuDOBGU7Ye2GHuUwEuNOXfAR4GthL2qU3An4EdwDdNWZf5vQV4Gfhd7J43A58CFhD25wHgQaBsyl4z7TwR6+PvAn806y0y7/Ue4GKznkb4KNACrAR2m/6fGutzW9e7ZdwVOy8c2MXDxra1sA5Y0OXpsjkb3xSL217N2fiyLk+X39HlLVDvxO3dZctT1t1lyvZO0A6r3V2efmiXp30EQ/bG9NPa9UxKGxt1XR1+TrytJ2NHbO1YUx8CvnqczAFgBfALQg8eATqA2cCfgJmmfhvwS0KHXAu8A/gpcK8pP5vw9r6F0ElvMm+gBtxOfZjy+th9H+a/1Av+QHCD9MAPRNdXmOu/Qxioe4zvBl5hnx4H/gK8E/gu8B3CWw8EdzkGPE5wmTYT3MG8o04B7iS8mYJ5ptxa8D3AbGAt8JgpWxwrW0BI1qzLsceU7yUMmB85WUYcXmf99+s2dm6ZazveJb3+r1b1Dc09LzTX8w3/wTEeO7mZsIGtrbnHM+paDDuO37NrgufUxj/nLuBzXZ4OdXna3eXpwjfYhleMfQe6PB3p8rSvy9PFKe3sijHbujxtNnVemKAtdwNPAzdNsP6UxxkJzPkAF0i2Rm8oGJ65wDJzbb+tl1OffCrANeZ3hZAD5Aiuwy3m3oGpvwu4KnZV59tHbxOCMcPvzfX6mKydesHrzPVSwhsQwttifw/G6i0itMEG5kuod5cWAReZ8kPUz12WEgKt4LmYvBC9+Qvgd+2Evsk+zSJELXaaZ+06jNiB/xEk5lPxnp5FwLOM//3X7HqOWXefue6aZD3ry2wn5EKfBc6l9ZAa7XrJEKzZlMCmYuE+AK9vUiP6gvj1Xq9P31dMrWEbXgceiFU9Oo77nT0JnquArzPVB4ilYmgc8hvAk6bsDNKD3z8I0cHxmJwrAMRkHuESgpHsw3/eXC8G3kQI5iFtv/z/IkLEtNjU301w0/KEJN/mGHah/yG4awOEgd1CiHq9QPpgzRG+tAFuiJW3Uw/+zwPupZ7TfC2lfkz/xZgkxHTnFQsHA2K3gvE9V5y/UB+QrVRGx9H+9NMsCRbKKTb9PofGOKILHJiEzm8Jg+S6CdadFphKcxBLG8EVsZ1lz34gviNvWbwNV0eyx6nH6nbnrYvV/RLwKUJgkKc+OHwgh/pJhNw9rFHrDa3Uj4TxtheBG6lHw+wnZl9C5ixm8l2RbK2pu9LIzqWeY8wjuFwbYnXPIEQ2N0V1Dzdoy17TTr5G8BK2mrKOFPnJ0Rdz/ORgbKwRjQbInMj52C7P2bB2WLVMzU21OHGOYTmTw/q0dZoTrPdWQs4HYXZo+9uSjzXkUGAD8BzpI2oDYXByFSF32BURTjP1yoQcJJ5I5qjvnjYRnP0Y8JC5z3pC7gFhp7cPd11Kew8RdnIb0W2NyNqMjh2o9nmWG97TxrYdhCiXpd3I4lG+Y7H2tZty6yK9UFBN1Yh6YnqC+u+kPvyE1PmWPf1eCCG/+j7waZrHKnkSrA5KG7CJaC3VZMcbMeVYQUiw4zlJLiXytHUWUP/xbeBldVhXa5GwkzbiU9QThpepB4BxLqV+vssRAs3VpKMZ+ATwNUL0q4W6OzBCcCHeE7OhQsgnZhF2+1tT7rcLuIm6S3eIENXbBvzRyE8nRLRWU9/dOxgtb3gu6gePPQCuMZf3jvOZG3EnYeADXFAsnCqfN+sTdZ2Jz0HirCBMkn2X1cLmD0N77eSZPY1oQtcgzPBvz59UbUhJIzl+ygYYaxPVDRn9VBugbJBDhPCazSfqDRyIeON5xGLTtr1lDcJxz1aCi7UzRWchJLnvIfT/PYR8rhX4oZHZlxNPtCvUT0Mu+p/qyoixPzL6QDSo2yn9A8VBwh+XtJjrFwkH7grprliOEPK7gpoT5RgfVQCadRPfKjqbYNtgk9zBn0JysFnXqJgir1IfHFZvsxPfasQD9V0+/hZWEhJzNZG7uUa6+9FIPwL4uJHlJ1GvkKhnx2JxOaP/iCrBY1HQMpcQoWvELKO7PcXmU6nPjW6xpFdPTnw0tPbQHDdTPdlUDfTHEuw4bx0/1V1ZTZo3aPQdDUuY8MRJxnGrVaPKRjmIH3CqA8R4o9VNGXVm3wBlK3WXK05cVoqVL5tQD6gSqTdASaxeX5PnTJKAuWa5t82+tMPR82yl/hJKRn6CsavtbGhUt6YcOQP2KepJrg3Fx5+7Qgg/2mAzfs7jVPRZV6g93E9G/0f0Xdwc1a0BVYKLWAXKqjNmHBu4DNWo5/A1a5LkYOOb675Szq4/UHzQUqjUpJUgbGvn8QOY42/1oziaqKe7GfFk24YXbTRsJSF3inMoqndJN1P/nKCY0sYK6guoThRfM7Yf5+TH5y21iC+mUqV+DG8s1W+eVplaDlFPrJ8kTM6PEKIqcxvoWZrRvF6BcP5qDyEw6iYE/xpxBfWEdgeh75cT+nYTYbAcoj6HH6ZxyDvO1YSzXA8RXuJHCUm6g21G9wlCDrIhpY4jHJq9DbiL8O+DdlN3KUo0dqXOAvYQNqOzqO+Q1gjNaVPKLmvv93EaHz203E34XnKO+nkvd0Z+9VuU9hMqA8WDAI5Kzp2LfOx6Ygn5tWm7/2lVMmr2AH2F6tCnc/WBfk1JVNW3VOOvvrkfEa8Svk6fCVwZK6tFdH1VNWjzGg08o/tEVLeqJkizdO1AEUqpnxZUUcnU3cxsxk6FRnlNLuZWrYvJC0Z3W0xejPTvicmrRn5rrL0hYLPSP/M4B9gy40jtqaFn2B8YGKkqv0JlMP65w2Jz3UcY7Huoj7Pt1HNgy+FoIYQ9Mf9+mnObYfM8hLngMHDMSJ8Enqwq+X5nXgmqZ5zrOirTZ9zqnWr3zKsBUNfhHdVLvKemJDUntFHLeYdOeQeAv2X/A5Xj0xXyVVB1VN+ynHP7XCUXQKmvGj4PbUZ2YMbwyiHQI9MHAQZF7hDAkFAHnJkXvNaXGysNBLWAo3SWcFunDxwYXjPjdQ9yt2d1gHBkzuTaUQb8oFTVVc4dUn2V0XOQV+sHZxX1HTpJPfG+d8bgDoCnl/lMVVGfQvVZVP8ZFUMBxfhgOUJ4oZfZ+Z5uGVZBdRPKfpTnRf8alDXAl3MH87kCb1W+P35qzmztZE+pbdOWHT3zl5ZztV28E9U1qFYRXoHqblQfxOn2sCMGCrXe0bEBAMMiFwnEL2d0FMvR/TkXEP1Dz9y2jtGZHWGnUroB1Tmo3uvseAgRQURuQbScJ1gMYTeLJLYfdZtJL0+jK3AdFXb7ZOQpvmB0aM2soPxJYO3YmFsEoKRfSlm3kJC1APxuzrGDPuXDOSWnk+xUCmARonLwzJXlAFCejf50QV8OVZeCFihp5QBe/Zf+P19O8vL3T0lm7EzomTt+vYyMDEv7/wD3CbMo0l4UYAAAACV0RVh0ZGF0ZTpjcmVhdGUAMjAyNS0wNC0xNVQwODoyNzoxMSswMDowMBKj9GgAAAAldEVYdGRhdGU6bW9kaWZ5ADIwMjUtMDQtMTVUMDg6Mjc6MTErMDA6MDBj/kzUAAAAAElFTkSuQmCC" alt="Bechtle Logo" class="logo">
    
    <div class="info-box">
        <p><strong>Server:</strong> {self.vcenter_server}</p>
        <p><strong>Benutzer:</strong> {self.username}</p>
        <p><strong>Erstellungsdatum:</strong> {timestamp}</p>
        <p><strong>Demo-Modus:</strong> {'Ja' if self.demo_mode else 'Nein'}</p>
    </div>

    <h1>VMware vSphere Infrastruktur-Bericht</h1>
"""
            
            # VMware Tools-Bericht
            if 'vmware_tools' in data and data['vmware_tools']:
                vmware_tools = data['vmware_tools']
                html_template += f"""
    <h2>VMware Tools Status ({len(vmware_tools)} VMs)</h2>
    <table>
        <tr>
            <th>VM-Name</th>
            <th>Power-Status</th>
            <th>Tools-Status</th>
            <th>Tools-Version</th>
            <th>Version-Status</th>
        </tr>
"""
                for vm in vmware_tools:
                    row_class = ''
                    if vm.get('status') == 'outdated':
                        row_class = 'warning'
                    
                    html_template += f"""
        <tr class="{row_class}">
            <td>{vm.get('name', 'N/A')}</td>
            <td>{vm.get('power_state', 'N/A')}</td>
            <td>{vm.get('tools_status', 'N/A')}</td>
            <td>{vm.get('tools_version', 'N/A')}</td>
            <td>{vm.get('tools_version_status', 'N/A')}</td>
        </tr>
"""
                
                html_template += "    </table>\n"
            
            # Snapshots-Bericht
            if 'snapshots' in data and data['snapshots']:
                snapshots = data['snapshots']
                html_template += f"""
    <h2>VM-Snapshots ({len(snapshots)} Snapshots)</h2>
    <table>
        <tr>
            <th>VM-Name</th>
            <th>Snapshot-Name</th>
            <th>Beschreibung</th>
            <th>Erstellungsdatum</th>
            <th>Alter</th>
            <th>Größe</th>
        </tr>
"""
                for snapshot in snapshots:
                    row_class = snapshot.get('status', '')
                    
                    html_template += f"""
        <tr class="{row_class}">
            <td>{snapshot.get('vm_name', 'N/A')}</td>
            <td>{snapshot.get('name', 'N/A')}</td>
            <td>{snapshot.get('description', 'N/A')}</td>
            <td>{snapshot.get('create_time', 'N/A')}</td>
            <td>{snapshot.get('age_text', 'N/A')}</td>
            <td>{snapshot.get('size_human', 'N/A')}</td>
        </tr>
"""
                
                html_template += "    </table>\n"
            
            # Verwaiste VMDKs-Bericht
            if 'orphaned_vmdks' in data and data['orphaned_vmdks']:
                orphaned_vmdks = data['orphaned_vmdks']
                html_template += f"""
    <h2>Verwaiste VMDK-Dateien ({len(orphaned_vmdks)} Dateien)</h2>
    <table>
        <tr>
            <th>Datastore</th>
            <th>Pfad</th>
            <th>Größe</th>
            <th>Zuletzt geändert</th>
        </tr>
"""
                for vmdk in orphaned_vmdks:
                    html_template += f"""
        <tr>
            <td>{vmdk.get('datastore', 'N/A')}</td>
            <td>{vmdk.get('path', 'N/A')}</td>
            <td>{vmdk.get('size_human', 'N/A')}</td>
            <td>{vmdk.get('modified', 'N/A')}</td>
        </tr>
"""
                
                html_template += "    </table>\n"
            
            # Footer
            html_template += """
    <div class="footer">
        <p>Erstellt mit VMware vSphere Reporter v29.0</p>
        <p>© 2025 Bechtle GmbH - Alle Rechte vorbehalten</p>
    </div>
</body>
</html>
"""
            
            # HTML-Datei schreiben
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_template)
            
            logger.info(f"HTML-Bericht erfolgreich erstellt: {output_file}")
            return output_file
            
        except Exception as e:
            logger.error(f"Fehler beim Generieren des HTML-Berichts: {str(e)}")
            raise ReportGenerationError(f"Fehler beim Generieren des HTML-Berichts: {str(e)}")
    
    @handle_vsphere_errors
    def generate_pdf_report(self, filename, data):
        """
        Generiert einen PDF-Bericht mit den gegebenen Daten.
        
        Args:
            filename: Dateiname für den Bericht
            data: Dictionary mit den Berichtsdaten
                
        Returns:
            str: Pfad zur generierten PDF-Datei
        """
        try:
            logger.info(f"Generiere PDF-Bericht: {filename}")
            
            # Prüfe, ob die PDF-Module verfügbar sind
            if not REPORT_MODULES_AVAILABLE:
                raise ReportGenerationError("PDF-Berichtsmodule nicht verfügbar. Installieren Sie reportlab.")
            
            # Stelle sicher, dass der Dateiname auf .pdf endet
            if not filename.lower().endswith('.pdf'):
                filename += '.pdf'
            
            # Vollständiger Pfad zur Ausgabedatei
            output_file = os.path.join(self.output_dir, filename)
            
            # PDF erstellen
            doc = SimpleDocTemplate(output_file, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # Eigene Stile definieren
            title_style = styles['Heading1']
            title_style.alignment = 1  # Zentriert
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.navy,
                spaceAfter=6
            )
            
            normal_style = ParagraphStyle(
                'Normal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6
            )
            
            header_style = ParagraphStyle(
                'Header',
                parent=styles['Normal'],
                fontSize=8,
                textColor=colors.gray
            )
            
            table_header_style = ParagraphStyle(
                'TableHeader',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.white,
                alignment=1
            )
            
            # Elemente für das PDF
            elements = []
            
            # Header mit Logo und Infos
            # Hier sollte ein Logo eingefügt werden, falls verfügbar
            
            # Titel
            elements.append(Paragraph("VMware vSphere Infrastruktur-Bericht", title_style))
            elements.append(Paragraph(f"Server: {self.vcenter_server}", normal_style))
            elements.append(Paragraph(f"Benutzer: {self.username}", normal_style))
            elements.append(Paragraph(f"Erstellungsdatum: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
            elements.append(Paragraph(f"Demo-Modus: {'Ja' if self.demo_mode else 'Nein'}", normal_style))
            
            elements.append(Spacer(1, 20))
            
            # VMware Tools-Bericht
            if 'vmware_tools' in data and data['vmware_tools']:
                vmware_tools = data['vmware_tools']
                elements.append(Paragraph(f"VMware Tools Status ({len(vmware_tools)} VMs)", subtitle_style))
                
                # Tabellendaten
                table_data = [["VM-Name", "Power-Status", "Tools-Status", "Tools-Version", "Version-Status"]]
                
                for vm in vmware_tools:
                    row = [
                        vm.get('name', 'N/A'),
                        vm.get('power_state', 'N/A'),
                        vm.get('tools_status', 'N/A'),
                        vm.get('tools_version', 'N/A'),
                        vm.get('tools_version_status', 'N/A')
                    ]
                    table_data.append(row)
                
                # Tabelle erstellen
                table = Table(table_data, repeatRows=1)
                
                # Tabellenformatierung
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.navy),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ])
                
                # Veraltete Tools hervorheben
                for i, vm in enumerate(vmware_tools):
                    if vm.get('status') == 'outdated':
                        table_style.add('BACKGROUND', (0, i+1), (-1, i+1), colors.lightyellow)
                
                table.setStyle(table_style)
                elements.append(table)
                elements.append(Spacer(1, 20))
            
            # Snapshots-Bericht
            if 'snapshots' in data and data['snapshots']:
                snapshots = data['snapshots']
                elements.append(Paragraph(f"VM-Snapshots ({len(snapshots)} Snapshots)", subtitle_style))
                
                # Tabellendaten
                table_data = [["VM-Name", "Snapshot-Name", "Erstellungsdatum", "Alter", "Größe"]]
                
                for snapshot in snapshots:
                    row = [
                        snapshot.get('vm_name', 'N/A'),
                        snapshot.get('name', 'N/A'),
                        snapshot.get('create_time', 'N/A'),
                        snapshot.get('age_text', 'N/A'),
                        snapshot.get('size_human', 'N/A')
                    ]
                    table_data.append(row)
                
                # Tabelle erstellen
                table = Table(table_data, repeatRows=1)
                
                # Tabellenformatierung
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.navy),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ])
                
                # Alte Snapshots hervorheben
                for i, snapshot in enumerate(snapshots):
                    if snapshot.get('status') == 'warning':
                        table_style.add('BACKGROUND', (0, i+1), (-1, i+1), colors.lightyellow)
                    elif snapshot.get('status') == 'danger':
                        table_style.add('BACKGROUND', (0, i+1), (-1, i+1), colors.lightpink)
                
                table.setStyle(table_style)
                elements.append(table)
                elements.append(Spacer(1, 20))
            
            # Verwaiste VMDKs-Bericht
            if 'orphaned_vmdks' in data and data['orphaned_vmdks']:
                orphaned_vmdks = data['orphaned_vmdks']
                elements.append(Paragraph(f"Verwaiste VMDK-Dateien ({len(orphaned_vmdks)} Dateien)", subtitle_style))
                
                # Tabellendaten
                table_data = [["Datastore", "Pfad", "Größe", "Zuletzt geändert"]]
                
                for vmdk in orphaned_vmdks:
                    row = [
                        vmdk.get('datastore', 'N/A'),
                        vmdk.get('path', 'N/A'),
                        vmdk.get('size_human', 'N/A'),
                        vmdk.get('modified', 'N/A')
                    ]
                    table_data.append(row)
                
                # Tabelle erstellen
                table = Table(table_data, repeatRows=1)
                
                # Tabellenformatierung
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.navy),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ])
                
                table.setStyle(table_style)
                elements.append(table)
                elements.append(Spacer(1, 20))
            
            # Footer
            elements.append(Paragraph("Erstellt mit VMware vSphere Reporter v29.0", header_style))
            elements.append(Paragraph("© 2025 Bechtle GmbH - Alle Rechte vorbehalten", header_style))
            
            # PDF erstellen
            doc.build(elements)
            
            logger.info(f"PDF-Bericht erfolgreich erstellt: {output_file}")
            return output_file
            
        except Exception as e:
            logger.error(f"Fehler beim Generieren des PDF-Berichts: {str(e)}")
            raise ReportGenerationError(f"Fehler beim Generieren des PDF-Berichts: {str(e)}")
    
    @handle_vsphere_errors
    def generate_docx_report(self, filename, data):
        """
        Generiert einen Word-Bericht mit den gegebenen Daten.
        
        Args:
            filename: Dateiname für den Bericht
            data: Dictionary mit den Berichtsdaten
                
        Returns:
            str: Pfad zur generierten DOCX-Datei
        """
        try:
            logger.info(f"Generiere DOCX-Bericht: {filename}")
            
            # Prüfe, ob die DOCX-Module verfügbar sind
            if not REPORT_MODULES_AVAILABLE:
                raise ReportGenerationError("DOCX-Berichtsmodule nicht verfügbar. Installieren Sie python-docx.")
            
            # Stelle sicher, dass der Dateiname auf .docx endet
            if not filename.lower().endswith('.docx'):
                filename += '.docx'
            
            # Vollständiger Pfad zur Ausgabedatei
            output_file = os.path.join(self.output_dir, filename)
            
            # DOCX erstellen
            document = docx.Document()
            
            # Seitenränder einstellen
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)
            
            # Titel und Infos
            title = document.add_heading("VMware vSphere Infrastruktur-Bericht", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            info_para = document.add_paragraph()
            info_para.add_run(f"Server: {self.vcenter_server}\n").bold = True
            info_para.add_run(f"Benutzer: {self.username}\n").bold = True
            info_para.add_run(f"Erstellungsdatum: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = True
            info_para.add_run(f"Demo-Modus: {'Ja' if self.demo_mode else 'Nein'}").bold = True
            
            document.add_paragraph()
            
            # VMware Tools-Bericht
            if 'vmware_tools' in data and data['vmware_tools']:
                vmware_tools = data['vmware_tools']
                document.add_heading(f"VMware Tools Status ({len(vmware_tools)} VMs)", 1)
                
                # Tabelle erstellen
                table = document.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Header-Zeile
                header_cells = table.rows[0].cells
                header_cells[0].text = "VM-Name"
                header_cells[1].text = "Power-Status"
                header_cells[2].text = "Tools-Status"
                header_cells[3].text = "Tools-Version"
                header_cells[4].text = "Version-Status"
                
                # Header formatieren
                for cell in header_cells:
                    cell_para = cell.paragraphs[0]
                    run = cell_para.runs[0]
                    run.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Hintergrundfarbe setzen (dunkelblau)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="00355E"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    # Textfarbe weiß
                    run.font.color.rgb = RGBColor(255, 255, 255)
                
                # Daten einfügen
                for vm in vmware_tools:
                    row_cells = table.add_row().cells
                    row_cells[0].text = vm.get('name', 'N/A')
                    row_cells[1].text = vm.get('power_state', 'N/A')
                    row_cells[2].text = vm.get('tools_status', 'N/A')
                    row_cells[3].text = vm.get('tools_version', 'N/A')
                    row_cells[4].text = vm.get('tools_version_status', 'N/A')
                    
                    # Veraltete Tools hervorheben
                    if vm.get('status') == 'outdated':
                        for cell in row_cells:
                            # Hintergrundfarbe setzen (hellgelb)
                            shading_elm = parse_xml(r'<w:shd {} w:fill="FFF9C4"/>'.format(nsdecls('w')))
                            cell._tc.get_or_add_tcPr().append(shading_elm)
                
                document.add_paragraph()
            
            # Snapshots-Bericht
            if 'snapshots' in data and data['snapshots']:
                snapshots = data['snapshots']
                document.add_heading(f"VM-Snapshots ({len(snapshots)} Snapshots)", 1)
                
                # Tabelle erstellen
                table = document.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Header-Zeile
                header_cells = table.rows[0].cells
                header_cells[0].text = "VM-Name"
                header_cells[1].text = "Snapshot-Name"
                header_cells[2].text = "Erstellungsdatum"
                header_cells[3].text = "Alter"
                header_cells[4].text = "Größe"
                
                # Header formatieren
                for cell in header_cells:
                    cell_para = cell.paragraphs[0]
                    run = cell_para.runs[0]
                    run.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Hintergrundfarbe setzen (dunkelblau)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="00355E"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    # Textfarbe weiß
                    run.font.color.rgb = RGBColor(255, 255, 255)
                
                # Daten einfügen
                for snapshot in snapshots:
                    row_cells = table.add_row().cells
                    row_cells[0].text = snapshot.get('vm_name', 'N/A')
                    row_cells[1].text = snapshot.get('name', 'N/A')
                    row_cells[2].text = snapshot.get('create_time', 'N/A')
                    row_cells[3].text = snapshot.get('age_text', 'N/A')
                    row_cells[4].text = snapshot.get('size_human', 'N/A')
                    
                    # Alte Snapshots hervorheben
                    status = snapshot.get('status')
                    if status:
                        for cell in row_cells:
                            if status == 'warning':
                                # Hintergrundfarbe setzen (hellgelb)
                                shading_elm = parse_xml(r'<w:shd {} w:fill="FFF9C4"/>'.format(nsdecls('w')))
                                cell._tc.get_or_add_tcPr().append(shading_elm)
                            elif status == 'danger':
                                # Hintergrundfarbe setzen (hellrot)
                                shading_elm = parse_xml(r'<w:shd {} w:fill="FFCDD2"/>'.format(nsdecls('w')))
                                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                document.add_paragraph()
            
            # Verwaiste VMDKs-Bericht
            if 'orphaned_vmdks' in data and data['orphaned_vmdks']:
                orphaned_vmdks = data['orphaned_vmdks']
                document.add_heading(f"Verwaiste VMDK-Dateien ({len(orphaned_vmdks)} Dateien)", 1)
                
                # Tabelle erstellen
                table = document.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Header-Zeile
                header_cells = table.rows[0].cells
                header_cells[0].text = "Datastore"
                header_cells[1].text = "Pfad"
                header_cells[2].text = "Größe"
                header_cells[3].text = "Zuletzt geändert"
                
                # Header formatieren
                for cell in header_cells:
                    cell_para = cell.paragraphs[0]
                    run = cell_para.runs[0]
                    run.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Hintergrundfarbe setzen (dunkelblau)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="00355E"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    # Textfarbe weiß
                    run.font.color.rgb = RGBColor(255, 255, 255)
                
                # Daten einfügen
                for vmdk in orphaned_vmdks:
                    row_cells = table.add_row().cells
                    row_cells[0].text = vmdk.get('datastore', 'N/A')
                    row_cells[1].text = vmdk.get('path', 'N/A')
                    row_cells[2].text = vmdk.get('size_human', 'N/A')
                    row_cells[3].text = vmdk.get('modified', 'N/A')
                
                document.add_paragraph()
            
            # Footer
            footer_para = document.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.add_run("Erstellt mit VMware vSphere Reporter v29.0\n").font.size = Pt(8)
            footer_para.add_run("© 2025 Bechtle GmbH - Alle Rechte vorbehalten").font.size = Pt(8)
            
            # Dokument speichern
            document.save(output_file)
            
            logger.info(f"DOCX-Bericht erfolgreich erstellt: {output_file}")
            return output_file
            
        except Exception as e:
            logger.error(f"Fehler beim Generieren des DOCX-Berichts: {str(e)}")
            raise ReportGenerationError(f"Fehler beim Generieren des DOCX-Berichts: {str(e)}")