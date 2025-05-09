"""
Bechtle vSphere Reporter v0.2 - Report Generator
Ermöglicht die Generierung und den Export von Berichten in verschiedenen Formaten (HTML, PDF, DOCX)

Diese Klasse verarbeitet die gesammelten Daten von vSphere und erstellt daraus strukturierte Berichte,
die in verschiedenen Formaten exportiert werden können.

© 2025 Bechtle GmbH - Alle Rechte vorbehalten
"""

import os
import logging
import time
from datetime import datetime
import tempfile
import json
import shutil
import uuid
from jinja2 import Environment, FileSystemLoader
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Logging einrichten
logger = logging.getLogger('vsphere_reporter.report_generator')

class ReportGenerator:
    """Generator für vSphere-Umgebungsberichte mit Export-Funktionen"""
    
    def __init__(self, data, client=None, demo_mode=False):
        """
        Initialisiert den Report-Generator
        
        Args:
            data (dict): Dictionary mit gesammelten vSphere-Daten
            client: VSphereClient-Instanz für Zusatzinformationen (optional)
            demo_mode (bool): Gibt an, ob Demo-Daten verwendet werden sollen
        """
        self.data = data or {}
        self.client = client
        self.demo_mode = demo_mode
        
        # Protokollieren der erhaltenen Daten für Diagnosezwecke
        logger.info("Report Generator initialisiert mit folgenden Daten:")
        logger.info(f"VMware Tools: {len(data.get('vmware_tools', []))} Einträge")
        logger.info(f"Snapshots: {len(data.get('snapshots', []))} Einträge")
        logger.info(f"Verwaiste VMDKs: {len(data.get('orphaned_vmdks', []))} Einträge")
        
        # Sicherstellen, dass alle erforderlichen Sektionen vorhanden sind
        if 'vmware_tools' not in self.data:
            logger.warning("Keine VMware Tools-Daten gefunden, erstelle leere Liste")
            self.data['vmware_tools'] = []
            
        if 'snapshots' not in self.data:
            logger.warning("Keine Snapshots-Daten gefunden, erstelle leere Liste")
            self.data['snapshots'] = []
            
        if 'orphaned_vmdks' not in self.data:
            logger.warning("Keine verwaisten VMDK-Daten gefunden, erstelle leere Liste")
            self.data['orphaned_vmdks'] = []
        
        # Erstellung eines Ordners für die Berichte
        self.reports_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'reports')
        os.makedirs(self.reports_dir, exist_ok=True)
        
        # Zeitstempel für eindeutige Benennungen
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Bechtle-Farbschema
        self.colors = {
            'primary': '#00355e',     # Dunkelblau
            'secondary': '#da6f1e',   # Orange
            'success': '#23a96a',     # Grün
            'light': '#f3f3f3',       # Hellgrau
            'dark': '#5a5a5a'         # Dunkelgrau
        }
        
        # Logo-Datei
        self.logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'logo.png')
    
    def generate_report(self, include_sections, export_formats):
        """
        Generiert einen Bericht basierend auf den ausgewählten Optionen
        
        Args:
            include_sections (dict): Dictionary mit den einzuschließenden Berichtsabschnitten
            export_formats (dict): Dictionary mit den zu generierenden Exportformaten
            
        Returns:
            dict: Pfade zu den generierten Berichten
        """
        logger.info(f"Starte Berichtgenerierung mit Abschnitten: {include_sections} und Formaten: {export_formats}")
        
        # Daten vorbereiten
        report_data = self._prepare_report_data(include_sections)
        
        # Liste für generierte Berichte
        generated_files = {}
        
        # HTML-Format exportieren
        if export_formats.get('html', False):
            html_path = self._generate_html_report(report_data)
            generated_files['html'] = html_path
            logger.info(f"HTML-Bericht erstellt: {html_path}")
        
        # PDF-Format exportieren
        if export_formats.get('pdf', False):
            pdf_path = self._generate_pdf_report(report_data)
            generated_files['pdf'] = pdf_path
            logger.info(f"PDF-Bericht erstellt: {pdf_path}")
        
        # DOCX-Format exportieren
        if export_formats.get('docx', False):
            docx_path = self._generate_docx_report(report_data)
            generated_files['docx'] = docx_path
            logger.info(f"DOCX-Bericht erstellt: {docx_path}")
        
        logger.info(f"Berichterstellung abgeschlossen. Generierte Dateien: {generated_files}")
        return generated_files
    
    def _prepare_report_data(self, include_sections):
        """
        Bereitet die Daten für den Bericht vor
        
        Args:
            include_sections (dict): Dictionary mit den einzuschließenden Berichtsabschnitten
            
        Returns:
            dict: Aufbereitete Daten für den Bericht
        """
        report_data = {
            'title': 'Bechtle vSphere Reporter - Infrastrukturbericht',
            'generated_at': datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            'demo_mode': self.demo_mode,
            'connection_info': {}
        }
        
        # Verbindungsinformationen hinzufügen
        if self.client and hasattr(self.client, 'connection_info'):
            report_data['connection_info'] = self.client.connection_info
        
        # Abschnitte basierend auf den ausgewählten Optionen hinzufügen
        if include_sections.get('vmware_tools', False):
            report_data['vmware_tools'] = self.data.get('vmware_tools', [])
        
        if include_sections.get('snapshots', False):
            report_data['snapshots'] = self.data.get('snapshots', [])
        
        if include_sections.get('orphaned_vmdks', False):
            report_data['orphaned_vmdks'] = self.data.get('orphaned_vmdks', [])
        
        # Zusammenfassung generieren
        summary = {
            'vm_tools_count': len(report_data.get('vmware_tools', [])),
            'snapshots_count': len(report_data.get('snapshots', [])),
            'orphaned_vmdks_count': len(report_data.get('orphaned_vmdks', [])),
            'orphaned_vmdks_size_gb': sum(vmdk.get('size_kb', 0) / (1024 * 1024) for vmdk in report_data.get('orphaned_vmdks', []))
        }
        report_data['summary'] = summary
        
        return report_data
    
    def _generate_html_report(self, report_data):
        """
        Generiert einen HTML-Bericht
        
        Args:
            report_data (dict): Aufbereitete Daten für den Bericht
            
        Returns:
            str: Pfad zur generierten HTML-Datei
        """
        # Jinja2-Umgebung einrichten
        env = Environment(loader=FileSystemLoader('templates'))
        template = env.get_template('report_template.html')
        
        # HTML-Datei erstellen
        filename = f"vsphere_report_{self.timestamp}.html"
        filepath = os.path.join(self.reports_dir, filename)
        
        # Template mit Daten rendern
        html_content = template.render(
            report=report_data,
            title=report_data['title'],
            colors=self.colors,
            demo_mode=self.demo_mode
        )
        
        # HTML in Datei schreiben
        with open(filepath, 'w', encoding='utf-8') as file:
            file.write(html_content)
        
        return filepath
    
    def _generate_pdf_report(self, report_data):
        """
        Generiert einen PDF-Bericht
        
        Args:
            report_data (dict): Aufbereitete Daten für den Bericht
            
        Returns:
            str: Pfad zur generierten PDF-Datei
        """
        # PDF-Datei erstellen
        filename = f"vsphere_report_{self.timestamp}.pdf"
        filepath = os.path.join(self.reports_dir, filename)
        
        # ReportLab-Dokument einrichten
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Styles definieren
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='BechteHeading1',
            fontName='Helvetica-Bold',
            fontSize=16,
            textColor=colors.HexColor(self.colors['primary']),
            spaceAfter=12
        ))
        styles.add(ParagraphStyle(
            name='BechteHeading2',
            fontName='Helvetica-Bold',
            fontSize=14,
            textColor=colors.HexColor(self.colors['primary']),
            spaceAfter=10
        ))
        styles.add(ParagraphStyle(
            name='BechteNormal',
            fontName='Helvetica',
            fontSize=10,
            spaceAfter=8
        ))
        styles.add(ParagraphStyle(
            name='BechtleDemoMode',
            fontName='Helvetica-Bold',
            fontSize=12,
            textColor=colors.HexColor(self.colors['secondary']),
            spaceAfter=8
        ))
        
        # Elemente für das PDF erstellen
        elements = []
        
        # Logo hinzufügen, wenn vorhanden
        if os.path.exists(self.logo_path):
            img_width = 150
            elements.append(Image(self.logo_path, width=img_width, height=img_width*0.5))
            elements.append(Spacer(1, 12))
        
        # Titel und Zeitstempel
        elements.append(Paragraph(report_data['title'], styles['BechteHeading1']))
        elements.append(Paragraph(f"Erstellt am: {report_data['generated_at']}", styles['BechteNormal']))
        
        # Demo-Modus-Hinweis
        if self.demo_mode:
            elements.append(Spacer(1, 12))
            elements.append(Paragraph("DEMO-MODUS - Alle Daten sind Beispieldaten", styles['BechtleDemoMode']))
        
        # Verbindungsinformationen
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Verbindungsinformationen", styles['BechteHeading2']))
        
        conn_info = report_data.get('connection_info', {})
        if conn_info:
            connection_data = [
                ["Server", conn_info.get('host', 'Unbekannt')],
                ["Benutzer", conn_info.get('username', 'Unbekannt')],
                ["SSL-Überprüfung", "Deaktiviert" if conn_info.get('disable_ssl_verification', False) else "Aktiviert"]
            ]
            
            # Tabellenstil definieren
            conn_table_style = TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor(self.colors['light'])),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor(self.colors['dark'])),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ])
            
            conn_table = Table(connection_data, colWidths=[100, 300])
            conn_table.setStyle(conn_table_style)
            elements.append(conn_table)
        
        # Zusammenfassung
        elements.append(Spacer(1, 20))
        elements.append(Paragraph("Zusammenfassung", styles['BechteHeading2']))
        
        summary = report_data.get('summary', {})
        summary_data = [
            ["Anzahl VMware Tools Einträge", str(summary.get('vm_tools_count', 0))],
            ["Anzahl Snapshots", str(summary.get('snapshots_count', 0))],
            ["Anzahl verwaister VMDKs", str(summary.get('orphaned_vmdks_count', 0))],
            ["Gesamtgröße verwaister VMDKs", f"{summary.get('orphaned_vmdks_size_gb', 0):.2f} GB"]
        ]
        
        # Tabellenstil für die Zusammenfassung
        summary_table_style = TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor(self.colors['light'])),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor(self.colors['dark'])),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
        ])
        
        summary_table = Table(summary_data, colWidths=[200, 200])
        summary_table.setStyle(summary_table_style)
        elements.append(summary_table)
        
        # Detailabschnitte basierend auf ausgewählten Optionen
        if 'vmware_tools' in report_data and report_data['vmware_tools']:
            elements.append(PageBreak())
            elements.append(Paragraph("VMware Tools Status", styles['BechteHeading2']))
            elements.append(Spacer(1, 10))
            
            # Tabellen-Header
            vmtools_data = [["VM-Name", "Tools-Status", "Version", "Update-Status"]]
            
            # Daten hinzufügen
            for vm in report_data['vmware_tools']:
                row = [
                    vm.get('vm_name', 'Unbekannt'),
                    vm.get('tools_status', 'Unbekannt'),
                    vm.get('tools_version', 'Unbekannt'),
                    vm.get('tools_update_status', 'Unbekannt')
                ]
                vmtools_data.append(row)
            
            # Tabellenstil mit Farbkodierung basierend auf Status
            vmtools_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['primary'])),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor(self.colors['light'])),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ])
            
            # Tabelle erstellen
            vmtools_table = Table(vmtools_data, repeatRows=1)
            vmtools_table.setStyle(vmtools_style)
            elements.append(vmtools_table)
            
            # Farbige Zeilen basierend auf dem Status
            for i, vm in enumerate(report_data['vmware_tools'], 1):
                status = vm.get('tools_status', '').lower()
                if 'not installed' in status:
                    bg_color = colors.HexColor('#ffcccc')  # Rot-ish für Fehler
                elif 'old' in status:
                    bg_color = colors.HexColor('#ffffcc')  # Gelb-ish für Warnung
                else:
                    bg_color = colors.white
                vmtools_table.setStyle(TableStyle([('BACKGROUND', (0, i), (-1, i), bg_color)]))
        
        # Snapshot-Daten
        if 'snapshots' in report_data and report_data['snapshots']:
            elements.append(PageBreak())
            elements.append(Paragraph("Snapshot-Übersicht", styles['BechteHeading2']))
            elements.append(Spacer(1, 10))
            
            # Tabellen-Header
            snapshots_data = [["VM-Name", "Snapshot-Name", "Erstellt am", "Alter (Tage)", "Größe"]]
            
            # Daten hinzufügen
            for snapshot in report_data['snapshots']:
                row = [
                    snapshot.get('vm_name', 'Unbekannt'),
                    snapshot.get('name', 'Unbekannt'),
                    snapshot.get('created_date', 'Unbekannt'),
                    str(snapshot.get('age_days', 0)),
                    f"{snapshot.get('size_gb', 0):.2f} GB"
                ]
                snapshots_data.append(row)
            
            # Tabellenstil mit Farbkodierung basierend auf Alter
            snapshots_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['primary'])),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor(self.colors['light'])),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ])
            
            # Tabelle erstellen
            snapshots_table = Table(snapshots_data, repeatRows=1)
            snapshots_table.setStyle(snapshots_style)
            elements.append(snapshots_table)
            
            # Farbige Zeilen basierend auf dem Alter
            for i, snapshot in enumerate(report_data['snapshots'], 1):
                age_days = snapshot.get('age_days', 0)
                if age_days > 30:
                    bg_color = colors.HexColor('#ffcccc')  # Rot-ish für alte Snapshots
                elif age_days > 7:
                    bg_color = colors.HexColor('#ffffcc')  # Gelb-ish für ältere Snapshots
                else:
                    bg_color = colors.HexColor('#ccffcc')  # Grün-ish für neue Snapshots
                snapshots_table.setStyle(TableStyle([('BACKGROUND', (0, i), (-1, i), bg_color)]))
        
        # Verwaiste VMDK-Daten
        if 'orphaned_vmdks' in report_data and report_data['orphaned_vmdks']:
            elements.append(PageBreak())
            elements.append(Paragraph("Verwaiste VMDK-Dateien", styles['BechteHeading2']))
            elements.append(Spacer(1, 10))
            
            # Erklärungstext
            elements.append(Paragraph(
                "Verwaiste VMDK-Dateien sind virtuelle Festplatten, die keiner virtuellen Maschine "
                "zugeordnet sind und potenziell gelöscht werden können, um Speicherplatz freizugeben.",
                styles['BechteNormal']
            ))
            elements.append(Spacer(1, 10))
            
            # Tabellen-Header
            vmdks_data = [["Dateiname", "Datastore", "Größe", "Änderungsdatum"]]
            
            # Daten hinzufügen
            for vmdk in report_data['orphaned_vmdks']:
                row = [
                    os.path.basename(vmdk.get('path', 'Unbekannt')),
                    vmdk.get('datastore', 'Unbekannt'),
                    f"{vmdk.get('size_kb', 0) / (1024 * 1024):.2f} GB",
                    vmdk.get('modification_time', 'Unbekannt')
                ]
                vmdks_data.append(row)
            
            # Tabellenstil
            vmdks_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['primary'])),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor(self.colors['light'])),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ])
            
            # Tabelle erstellen
            vmdks_table = Table(vmdks_data, repeatRows=1)
            vmdks_table.setStyle(vmdks_style)
            elements.append(vmdks_table)
        
        # PDF erstellen
        doc.build(elements)
        
        return filepath
    
    def _generate_docx_report(self, report_data):
        """
        Generiert einen DOCX-Bericht
        
        Args:
            report_data (dict): Aufbereitete Daten für den Bericht
            
        Returns:
            str: Pfad zur generierten DOCX-Datei
        """
        # DOCX-Datei erstellen
        filename = f"vsphere_report_{self.timestamp}.docx"
        filepath = os.path.join(self.reports_dir, filename)
        
        # Neues Word-Dokument erstellen
        doc = Document()
        
        # Dokumenteigenschaften setzen
        doc.core_properties.title = "VMware vSphere Infrastrukturbericht"
        doc.core_properties.author = "Bechtle vSphere Reporter"
        
        # Bechtle-Farben als RGB für DOCX
        bechtle_blue = RGBColor(0, 53, 94)      # #00355e
        bechtle_orange = RGBColor(218, 111, 30)  # #da6f1e
        bechtle_green = RGBColor(35, 169, 106)   # #23a96a
        
        # Titel hinzufügen
        title = doc.add_heading(report_data['title'], level=1)
        for run in title.runs:
            run.font.color.rgb = bechtle_blue
        
        # Zeitstempel hinzufügen
        doc.add_paragraph(f"Erstellt am: {report_data['generated_at']}")
        
        # Demo-Modus-Hinweis
        if self.demo_mode:
            demo_para = doc.add_paragraph()
            demo_run = demo_para.add_run("DEMO-MODUS - Alle Daten sind Beispieldaten")
            demo_run.font.bold = True
            demo_run.font.color.rgb = bechtle_orange
        
        # Verbindungsinformationen
        doc.add_heading("Verbindungsinformationen", level=2)
        
        conn_info = report_data.get('connection_info', {})
        if conn_info:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Eigenschaft"
            hdr_cells[1].text = "Wert"
            
            # Header formatieren
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Daten hinzufügen
            row_cells = table.add_row().cells
            row_cells[0].text = "Server"
            row_cells[1].text = conn_info.get('host', 'Unbekannt')
            
            row_cells = table.add_row().cells
            row_cells[0].text = "Benutzer"
            row_cells[1].text = conn_info.get('username', 'Unbekannt')
            
            row_cells = table.add_row().cells
            row_cells[0].text = "SSL-Überprüfung"
            row_cells[1].text = "Deaktiviert" if conn_info.get('disable_ssl_verification', False) else "Aktiviert"
        
        # Zusammenfassung
        doc.add_heading("Zusammenfassung", level=2)
        summary = report_data.get('summary', {})
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metrik"
        hdr_cells[1].text = "Wert"
        
        # Header formatieren
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Zusammenfassungsdaten
        row_cells = table.add_row().cells
        row_cells[0].text = "Anzahl VMware Tools Einträge"
        row_cells[1].text = str(summary.get('vm_tools_count', 0))
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Anzahl Snapshots"
        row_cells[1].text = str(summary.get('snapshots_count', 0))
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Anzahl verwaister VMDKs"
        row_cells[1].text = str(summary.get('orphaned_vmdks_count', 0))
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Gesamtgröße verwaister VMDKs"
        row_cells[1].text = f"{summary.get('orphaned_vmdks_size_gb', 0):.2f} GB"
        
        # Detailabschnitte basierend auf ausgewählten Optionen
        # VMware Tools-Abschnitt
        if 'vmware_tools' in report_data and report_data['vmware_tools']:
            doc.add_page_break()
            doc.add_heading("VMware Tools Status", level=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "VM-Name"
            hdr_cells[1].text = "Tools-Status"
            hdr_cells[2].text = "Version"
            hdr_cells[3].text = "Update-Status"
            
            # Header formatieren
            for cell in hdr_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Daten hinzufügen
            for vm in report_data['vmware_tools']:
                row_cells = table.add_row().cells
                row_cells[0].text = vm.get('vm_name', 'Unbekannt')
                row_cells[1].text = vm.get('tools_status', 'Unbekannt')
                row_cells[2].text = vm.get('tools_version', 'Unbekannt')
                row_cells[3].text = vm.get('tools_update_status', 'Unbekannt')
        
        # Snapshot-Abschnitt
        if 'snapshots' in report_data and report_data['snapshots']:
            doc.add_page_break()
            doc.add_heading("Snapshot-Übersicht", level=2)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "VM-Name"
            hdr_cells[1].text = "Snapshot-Name"
            hdr_cells[2].text = "Erstellt am"
            hdr_cells[3].text = "Alter (Tage)"
            hdr_cells[4].text = "Größe"
            
            # Header formatieren
            for cell in hdr_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Daten hinzufügen
            for snapshot in report_data['snapshots']:
                row_cells = table.add_row().cells
                row_cells[0].text = snapshot.get('vm_name', 'Unbekannt')
                row_cells[1].text = snapshot.get('name', 'Unbekannt')
                row_cells[2].text = snapshot.get('created_date', 'Unbekannt')
                row_cells[3].text = str(snapshot.get('age_days', 0))
                row_cells[4].text = f"{snapshot.get('size_gb', 0):.2f} GB"
        
        # Verwaiste VMDK-Abschnitt
        if 'orphaned_vmdks' in report_data and report_data['orphaned_vmdks']:
            doc.add_page_break()
            doc.add_heading("Verwaiste VMDK-Dateien", level=2)
            
            # Erklärungstext
            doc.add_paragraph(
                "Verwaiste VMDK-Dateien sind virtuelle Festplatten, die keiner virtuellen "
                "Maschine zugeordnet sind und potenziell gelöscht werden können, um Speicherplatz freizugeben."
            )
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Dateiname"
            hdr_cells[1].text = "Datastore"
            hdr_cells[2].text = "Größe"
            hdr_cells[3].text = "Änderungsdatum"
            
            # Header formatieren
            for cell in hdr_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Daten hinzufügen
            for vmdk in report_data['orphaned_vmdks']:
                row_cells = table.add_row().cells
                row_cells[0].text = os.path.basename(vmdk.get('path', 'Unbekannt'))
                row_cells[1].text = vmdk.get('datastore', 'Unbekannt')
                row_cells[2].text = f"{vmdk.get('size_kb', 0) / (1024 * 1024):.2f} GB"
                row_cells[3].text = vmdk.get('modification_time', 'Unbekannt')
        
        # Dokument speichern
        doc.save(filepath)
        
        return filepath