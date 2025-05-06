#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VMware vSphere Reporter v29.0
Report Generator Module

Final Fixed Version 10 - Simplified Structure
Copyright (c) 2025 Bechtle GmbH
"""

import os
import logging
import datetime
import time
from pathlib import Path
import jinja2
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Import custom modules
import demo_data
from vsphere_client import VSphereClient, VSphereConnectionError

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generator for VMware vSphere reports in various formats"""
    
    def __init__(self, demo_mode, vsphere_host=None, vsphere_user=None, vsphere_password=None, ignore_ssl=False):
        """Initialize the report generator
        
        Args:
            demo_mode (bool): Whether to use demo data or real vCenter connection
            vsphere_host (str): vCenter hostname or IP (only used if demo_mode is False)
            vsphere_user (str): vCenter username (only used if demo_mode is False)
            vsphere_password (str): vCenter password (only used if demo_mode is False)
            ignore_ssl (bool): Whether to ignore SSL certificate validation (only used if demo_mode is False)
        """
        self.demo_mode = demo_mode
        self.vsphere_host = vsphere_host
        self.vsphere_user = vsphere_user
        self.vsphere_password = vsphere_password
        self.ignore_ssl = ignore_ssl
        self.client = None
        
        # Bechtle corporate colors
        self.colors = {
            'dark_blue': '#00355e',
            'orange': '#da6f1e',
            'green': '#23a96a',
            'light_gray': '#f3f3f3',
            'dark_gray': '#5a5a5a'
        }
        
        # Initialize Jinja2 environment
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader('templates'),
            autoescape=jinja2.select_autoescape(['html', 'xml'])
        )
    
    def _get_data(self):
        """Get data for the report
        
        This method either uses demo data or connects to a real vCenter server
        to collect the necessary information.
        
        Returns:
            dict: Dictionary containing all data for the report
        """
        data = {}
        
        if self.demo_mode:
            # Use demo data
            data['total_vms'] = demo_data.get_total_vms()
            data['total_hosts'] = demo_data.get_total_hosts()
            data['total_datastores'] = demo_data.get_total_datastores()
            data['total_clusters'] = demo_data.get_total_clusters()
            data['vms_by_power_state'] = demo_data.get_vms_by_power_state()
            data['vmware_tools'] = demo_data.get_vmware_tools_data()
            data['snapshots'] = demo_data.get_snapshots_data()
            data['orphaned_vmdks'] = demo_data.get_orphaned_vmdks_data()
            data['connection_info'] = {
                'host': 'demo.vcenter.local',
                'user': 'demo@vsphere.local',
                'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
        else:
            # Connect to real vCenter
            try:
                self.client = VSphereClient(
                    self.vsphere_host,
                    self.vsphere_user,
                    self.vsphere_password,
                    self.ignore_ssl
                )
                self.client.connect()
                
                # Collect real data
                data['total_vms'] = self.client.get_total_vms()
                data['total_hosts'] = self.client.get_total_hosts()
                data['total_datastores'] = self.client.get_total_datastores()
                data['total_clusters'] = self.client.get_total_clusters()
                data['vms_by_power_state'] = self.client.get_vms_by_power_state()
                data['vmware_tools'] = self.client.get_vmware_tools_data()
                data['snapshots'] = self.client.get_snapshots_data()
                data['orphaned_vmdks'] = self.client.get_orphaned_vmdks_data()
                data['connection_info'] = {
                    'host': self.vsphere_host,
                    'user': self.vsphere_user,
                    'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                }
            except Exception as e:
                logger.error(f"Fehler beim Sammeln der vCenter-Daten: {str(e)}")
                raise
            finally:
                if self.client:
                    self.client.disconnect()
        
        # Add processing information
        # Sort VMware Tools data (outdated first)
        data['vmware_tools'].sort(key=lambda x: (
            x['version'] == 'Not installed', 
            x['status'] == 'Outdated', 
            x['status'] == 'Up-to-date'
        ))
        
        # Sort snapshots by age (oldest first)
        now = datetime.datetime.now()
        for snapshot in data['snapshots']:
            creation_time = snapshot.get('creation_time')
            if creation_time:
                # Make sure both datetimes are naive (no timezone info)
                if hasattr(creation_time, 'tzinfo') and creation_time.tzinfo is not None:
                    creation_time = creation_time.replace(tzinfo=None)
                    
                age_seconds = (now - creation_time).total_seconds()
                snapshot['age_days'] = age_seconds / (24 * 60 * 60)
                snapshot['age_human'] = self._format_timespan(age_seconds)
        
        data['snapshots'].sort(key=lambda x: x.get('creation_time', now))
        
        # Sort orphaned VMDKs by size (largest first)
        for vmdk in data['orphaned_vmdks']:
            size_bytes = vmdk.get('size_bytes', 0)
            vmdk['size_human'] = self._format_size(size_bytes)
        
        data['orphaned_vmdks'].sort(key=lambda x: x.get('size_bytes', 0), reverse=True)
        
        return data
    
    def _format_timespan(self, seconds):
        """Format a timespan in seconds to a human-readable string
        
        Args:
            seconds (float): Timespan in seconds
            
        Returns:
            str: Human-readable timespan
        """
        if seconds < 60:
            return f"{int(seconds)} seconds"
        elif seconds < 3600:
            return f"{int(seconds / 60)} minutes"
        elif seconds < 86400:
            return f"{int(seconds / 3600)} hours"
        else:
            days = int(seconds / 86400)
            return f"{days} days"
    
    def _format_size(self, size_bytes):
        """Format a size in bytes to a human-readable string
        
        Args:
            size_bytes (int): Size in bytes
            
        Returns:
            str: Human-readable size
        """
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.2f} KB"
        elif size_bytes < 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.2f} MB"
        else:
            return f"{size_bytes / (1024 * 1024 * 1024):.2f} GB"
    
    def generate_report(self, output_dir, formats, options):
        """Generate reports in the specified formats
        
        Args:
            output_dir (str): Directory to save the reports
            formats (list): List of formats to generate ('html', 'pdf', 'docx')
            options (dict): Dictionary of report sections to include
            
        Returns:
            dict: Dictionary of generated report files
        """
        logger.info(f"Generiere Bericht in Formaten: {', '.join(formats)}")
        
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        # Get report data
        try:
            data = self._get_data()
        except Exception as e:
            logger.error(f"Fehler beim Sammeln der Daten für den Bericht: {str(e)}")
            raise
        
        # Generate reports in selected formats
        generated_files = {}
        
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if 'html' in formats:
            html_file = os.path.join(output_dir, f"vsphere_report_{timestamp}.html")
            self._generate_html_report(html_file, data, options)
            generated_files['html'] = os.path.basename(html_file)
        
        if 'pdf' in formats:
            pdf_file = os.path.join(output_dir, f"vsphere_report_{timestamp}.pdf")
            self._generate_pdf_report(pdf_file, data, options)
            generated_files['pdf'] = os.path.basename(pdf_file)
        
        if 'docx' in formats:
            docx_file = os.path.join(output_dir, f"vsphere_report_{timestamp}.docx")
            self._generate_docx_report(docx_file, data, options)
            generated_files['docx'] = os.path.basename(docx_file)
        
        logger.info(f"Bericht erfolgreich generiert: {', '.join(generated_files.values())}")
        return generated_files
    
    def _generate_html_report(self, output_file, data, options):
        """Generate HTML report
        
        Args:
            output_file (str): Path to save the HTML report
            data (dict): Report data
            options (dict): Report options
        """
        try:
            # Create Jinja2 template context
            context = {
                'title': 'VMware vSphere Reporter',
                'generation_time': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'connection_info': data['connection_info'],
                'options': options,
                'colors': self.colors,
                'demo_mode': self.demo_mode
            }
            
            # Add data based on selected options
            if options.get('vmware_tools', False):
                context['vmware_tools'] = data['vmware_tools']
            
            if options.get('snapshots', False):
                context['snapshots'] = data['snapshots']
            
            if options.get('orphaned_vmdks', False):
                context['orphaned_vmdks'] = data['orphaned_vmdks']
            
            if options.get('vm_hardware', False):
                context['vm_hardware'] = data.get('vm_hardware', [])
            
            if options.get('datastores', False):
                context['datastores'] = data.get('datastores', [])
            
            if options.get('hosts', False):
                context['hosts'] = data.get('hosts', [])
            
            if options.get('clusters', False):
                context['clusters'] = data.get('clusters', [])
            
            # Summary data (always included)
            context['summary'] = {
                'total_vms': data['total_vms'],
                'total_hosts': data['total_hosts'],
                'total_datastores': data['total_datastores'],
                'total_clusters': data['total_clusters'],
                'vms_by_power_state': data['vms_by_power_state']
            }
            
            # Render the template
            template_file = os.path.join('templates', 'report_template.html')
            if os.path.exists(template_file):
                with open(template_file, 'r', encoding='utf-8') as f:
                    template_content = f.read()
                template = jinja2.Template(template_content)
                html_content = template.render(**context)
            else:
                # Fallback to creating a simple HTML file
                html_content = self._generate_fallback_html(context)
            
            # Write HTML file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"HTML-Bericht gespeichert: {output_file}")
        
        except Exception as e:
            logger.error(f"Fehler bei der Generierung des HTML-Berichts: {str(e)}")
            raise
    
    def _generate_fallback_html(self, context):
        """Generate a simple HTML report without using Jinja2 templates
        
        Args:
            context (dict): Template context
            
        Returns:
            str: HTML content
        """
        # Create a basic HTML structure
        html = [
            "<!DOCTYPE html>",
            "<html lang='de'>",
            "<head>",
            "    <meta charset='UTF-8'>",
            "    <meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            f"    <title>{context['title']}</title>",
            "    <style>",
            f"        body {{ font-family: Arial, sans-serif; color: {self.colors['dark_gray']}; margin: 0; padding: 0; }}",
            f"        header {{ background-color: {self.colors['dark_blue']}; color: white; padding: 20px; }}",
            "        h1, h2, h3 { margin-top: 20px; }",
            "        .container { max-width: 1200px; margin: 0 auto; padding: 20px; }",
            "        table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }",
            "        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "        th { background-color: #f3f3f3; }",
            f"        .section {{ background-color: {self.colors['light_gray']}; margin-bottom: 20px; padding: 15px; border-radius: 5px; }}",
            f"        .warning {{ color: {self.colors['orange']}; }}",
            f"        .success {{ color: {self.colors['green']}; }}",
            "        footer { margin-top: 30px; font-size: 0.8em; color: #777; text-align: center; }",
            "    </style>",
            "</head>",
            "<body>",
            "    <header>",
            f"        <h1>{context['title']}</h1>",
            "    </header>",
            "    <div class='container'>"
        ]
        
        # Add summary section
        html.append("<div class='section'>")
        html.append("<h2>Zusammenfassung</h2>")
        html.append("<p>Übersicht der vSphere-Umgebung:</p>")
        html.append("<ul>")
        html.append(f"<li>VMs: {context['summary']['total_vms']}</li>")
        html.append(f"<li>Hosts: {context['summary']['total_hosts']}</li>")
        html.append(f"<li>Datastores: {context['summary']['total_datastores']}</li>")
        html.append(f"<li>Cluster: {context['summary']['total_clusters']}</li>")
        html.append("</ul>")
        html.append("</div>")
        
        # Add VMware Tools section if selected
        if context['options'].get('vmware_tools', False) and 'vmware_tools' in context:
            html.append("<div class='section'>")
            html.append("<h2>VMware Tools Status</h2>")
            html.append("<table>")
            html.append("<tr><th>VM Name</th><th>Status</th><th>Version</th><th>Betriebssystem</th><th>Power State</th></tr>")
            
            for tool in context['vmware_tools']:
                status_class = ""
                if tool['status'] == "Outdated":
                    status_class = "warning"
                elif tool['status'] == "Up-to-date":
                    status_class = "success"
                
                html.append("<tr>")
                html.append(f"<td>{tool['vm_name']}</td>")
                html.append(f"<td class='{status_class}'>{tool['status']}</td>")
                html.append(f"<td>{tool['version']}</td>")
                html.append(f"<td>{tool['os']}</td>")
                html.append(f"<td>{tool['power_state']}</td>")
                html.append("</tr>")
            
            html.append("</table>")
            html.append("</div>")
        
        # Add Snapshots section if selected
        if context['options'].get('snapshots', False) and 'snapshots' in context:
            html.append("<div class='section'>")
            html.append("<h2>Snapshots</h2>")
            html.append("<table>")
            html.append("<tr><th>VM Name</th><th>Snapshot Name</th><th>Alter</th><th>Größe</th><th>Beschreibung</th></tr>")
            
            for snapshot in context['snapshots']:
                age_class = ""
                if snapshot.get('age_days', 0) > 30:
                    age_class = "warning"
                
                html.append("<tr>")
                html.append(f"<td>{snapshot['vm_name']}</td>")
                html.append(f"<td>{snapshot['name']}</td>")
                html.append(f"<td class='{age_class}'>{snapshot.get('age_human', 'Unbekannt')}</td>")
                html.append(f"<td>{snapshot.get('size_human', 'Unbekannt')}</td>")
                html.append(f"<td>{snapshot.get('description', '')}</td>")
                html.append("</tr>")
            
            html.append("</table>")
            html.append("</div>")
        
        # Add Orphaned VMDKs section if selected
        if context['options'].get('orphaned_vmdks', False) and 'orphaned_vmdks' in context:
            html.append("<div class='section'>")
            html.append("<h2>Verwaiste VMDK-Dateien</h2>")
            html.append("<table>")
            html.append("<tr><th>Dateiname</th><th>Datastore</th><th>Größe</th><th>Empfohlene Aktion</th></tr>")
            
            for vmdk in context['orphaned_vmdks']:
                html.append("<tr>")
                html.append(f"<td>{vmdk['name']}</td>")
                html.append(f"<td>{vmdk['datastore']}</td>")
                html.append(f"<td>{vmdk.get('size_human', 'Unbekannt')}</td>")
                html.append(f"<td>{vmdk.get('recommended_action', 'Manuelle Überprüfung erforderlich')}</td>")
                html.append("</tr>")
            
            html.append("</table>")
            html.append("</div>")
        
        # Add footer
        html.append("    </div>")
        html.append("    <footer>")
        html.append(f"        <p>Generiert am {context['generation_time']} | VMware vSphere Reporter v29.0 | &copy; 2025 Bechtle GmbH</p>")
        if context['demo_mode']:
            html.append("        <p><strong>Demo-Modus: Die angezeigten Daten sind Beispieldaten und keine echten vSphere-Daten.</strong></p>")
        html.append("    </footer>")
        html.append("</body>")
        html.append("</html>")
        
        return "\n".join(html)
    
    def _generate_pdf_report(self, output_file, data, options):
        """Generate PDF report using ReportLab
        
        Args:
            output_file (str): Path to save the PDF report
            data (dict): Report data
            options (dict): Report options
        """
        try:
            # Create a PDF document
            doc = SimpleDocTemplate(
                output_file,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Create styles
            styles = getSampleStyleSheet()
            title_style = styles['Title']
            heading1_style = styles['Heading1']
            heading2_style = styles['Heading2']
            normal_style = styles['Normal']
            
            # Custom styles for Bechtle corporate design
            title_style.textColor = colors.HexColor(self.colors['dark_blue'])
            heading1_style.textColor = colors.HexColor(self.colors['dark_blue'])
            heading2_style.textColor = colors.HexColor(self.colors['dark_blue'])
            
            # Define warning and success styles
            warning_style = ParagraphStyle(
                'Warning',
                parent=normal_style,
                textColor=colors.HexColor(self.colors['orange'])
            )
            success_style = ParagraphStyle(
                'Success',
                parent=normal_style,
                textColor=colors.HexColor(self.colors['green'])
            )
            
            # Build the document
            elements = []
            
            # Title
            elements.append(Paragraph("VMware vSphere Report", title_style))
            elements.append(Spacer(1, 12))
            
            # Connection info
            elements.append(Paragraph("Verbindungsinformationen:", heading2_style))
            elements.append(Paragraph(f"vCenter: {data['connection_info']['host']}", normal_style))
            elements.append(Paragraph(f"Benutzer: {data['connection_info']['user']}", normal_style))
            elements.append(Paragraph(f"Generiert am: {data['connection_info']['timestamp']}", normal_style))
            
            if self.demo_mode:
                elements.append(Paragraph("DEMO-MODUS: Die angezeigten Daten sind Beispieldaten.", warning_style))
            
            elements.append(Spacer(1, 12))
            
            # Summary
            elements.append(Paragraph("Zusammenfassung", heading1_style))
            
            summary_data = [
                ['Kategorie', 'Anzahl'],
                ['VMs', str(data['total_vms'])],
                ['Hosts', str(data['total_hosts'])],
                ['Datastores', str(data['total_datastores'])],
                ['Cluster', str(data['total_clusters'])]
            ]
            
            summary_table = Table(summary_data, colWidths=[200, 100])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['light_gray'])),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(self.colors['dark_blue'])),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            
            elements.append(summary_table)
            elements.append(Spacer(1, 24))
            
            # Add VMware Tools section if selected
            if options.get('vmware_tools', False):
                elements.append(Paragraph("VMware Tools Status", heading1_style))
                
                if data['vmware_tools']:
                    # Table header
                    vmware_tools_data = [['VM Name', 'Status', 'Version', 'Betriebssystem', 'Power State']]
                    
                    # Table data
                    for tool in data['vmware_tools']:
                        vmware_tools_data.append([
                            tool['vm_name'],
                            tool['status'],
                            tool['version'],
                            tool['os'],
                            tool['power_state']
                        ])
                    
                    # Create table
                    vmware_tools_table = Table(vmware_tools_data, colWidths=[100, 70, 70, 150, 80])
                    vmware_tools_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['light_gray'])),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(self.colors['dark_blue'])),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    
                    elements.append(vmware_tools_table)
                else:
                    elements.append(Paragraph("Keine VMware Tools-Daten verfügbar.", normal_style))
                
                elements.append(Spacer(1, 24))
            
            # Add Snapshots section if selected
            if options.get('snapshots', False):
                elements.append(Paragraph("Snapshots", heading1_style))
                
                if data['snapshots']:
                    # Table header
                    snapshots_data = [['VM Name', 'Snapshot Name', 'Alter', 'Größe', 'Beschreibung']]
                    
                    # Table data
                    for snapshot in data['snapshots']:
                        snapshots_data.append([
                            snapshot['vm_name'],
                            snapshot['name'],
                            snapshot.get('age_human', 'Unbekannt'),
                            str(snapshot.get('size_mb', 'Unbekannt')) + ' MB' if snapshot.get('size_mb') else 'Unbekannt',
                            snapshot.get('description', '')
                        ])
                    
                    # Create table
                    snapshots_table = Table(snapshots_data, colWidths=[100, 100, 70, 70, 130])
                    snapshots_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['light_gray'])),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(self.colors['dark_blue'])),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    
                    elements.append(snapshots_table)
                else:
                    elements.append(Paragraph("Keine Snapshot-Daten verfügbar.", normal_style))
                
                elements.append(Spacer(1, 24))
            
            # Add Orphaned VMDKs section if selected
            if options.get('orphaned_vmdks', False):
                elements.append(Paragraph("Verwaiste VMDK-Dateien", heading1_style))
                
                if data['orphaned_vmdks']:
                    # Table header
                    orphaned_vmdks_data = [['Dateiname', 'Datastore', 'Größe', 'Empfohlene Aktion']]
                    
                    # Table data
                    for vmdk in data['orphaned_vmdks']:
                        orphaned_vmdks_data.append([
                            vmdk['name'],
                            vmdk['datastore'],
                            vmdk.get('size_human', 'Unbekannt'),
                            vmdk.get('recommended_action', 'Manuelle Überprüfung erforderlich')
                        ])
                    
                    # Create table
                    orphaned_vmdks_table = Table(orphaned_vmdks_data, colWidths=[150, 100, 70, 150])
                    orphaned_vmdks_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['light_gray'])),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor(self.colors['dark_blue'])),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    
                    elements.append(orphaned_vmdks_table)
                else:
                    elements.append(Paragraph("Keine verwaisten VMDK-Dateien gefunden.", normal_style))
                
                elements.append(Spacer(1, 24))
            
            # Footer
            footer_text = f"Generiert mit VMware vSphere Reporter v29.0 | © 2025 Bechtle GmbH"
            elements.append(Paragraph(footer_text, normal_style))
            
            # Build PDF
            doc.build(elements)
            
            logger.info(f"PDF-Bericht gespeichert: {output_file}")
        
        except Exception as e:
            logger.error(f"Fehler bei der Generierung des PDF-Berichts: {str(e)}")
            raise
    
    def _generate_docx_report(self, output_file, data, options):
        """Generate Word DOCX report using python-docx
        
        Args:
            output_file (str): Path to save the DOCX report
            data (dict): Report data
            options (dict): Report options
        """
        try:
            # Create a new Document
            doc = Document()
            
            # Add title
            doc.add_heading('VMware vSphere Report', 0)
            
            # Add connection info
            doc.add_heading('Verbindungsinformationen', 1)
            doc.add_paragraph(f"vCenter: {data['connection_info']['host']}")
            doc.add_paragraph(f"Benutzer: {data['connection_info']['user']}")
            doc.add_paragraph(f"Generiert am: {data['connection_info']['timestamp']}")
            
            if self.demo_mode:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("DEMO-MODUS: Die angezeigten Daten sind Beispieldaten.")
                run.font.color.rgb = RGBColor.from_string(self.colors['orange'][1:])
                run.bold = True
            
            # Add summary
            doc.add_heading('Zusammenfassung', 1)
            
            # Create summary table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add table header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Kategorie'
            header_cells[1].text = 'Anzahl'
            
            # Add table data
            categories = [
                ('VMs', data['total_vms']),
                ('Hosts', data['total_hosts']),
                ('Datastores', data['total_datastores']),
                ('Cluster', data['total_clusters'])
            ]
            
            for category, count in categories:
                row_cells = table.add_row().cells
                row_cells[0].text = category
                row_cells[1].text = str(count)
            
            doc.add_paragraph()
            
            # Add VMware Tools section if selected
            if options.get('vmware_tools', False):
                doc.add_heading('VMware Tools Status', 1)
                
                if data['vmware_tools']:
                    # Create table
                    table = doc.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    
                    # Add table header
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'VM Name'
                    header_cells[1].text = 'Status'
                    header_cells[2].text = 'Version'
                    header_cells[3].text = 'Betriebssystem'
                    header_cells[4].text = 'Power State'
                    
                    # Add table data
                    for tool in data['vmware_tools']:
                        row_cells = table.add_row().cells
                        row_cells[0].text = tool['vm_name']
                        row_cells[1].text = tool['status']
                        row_cells[2].text = tool['version']
                        row_cells[3].text = tool['os']
                        row_cells[4].text = tool['power_state']
                        
                        # Highlight outdated tools
                        if tool['status'] == 'Outdated':
                            for paragraph in row_cells[1].paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor.from_string(self.colors['orange'][1:])
                else:
                    doc.add_paragraph("Keine VMware Tools-Daten verfügbar.")
                
                doc.add_paragraph()
            
            # Add Snapshots section if selected
            if options.get('snapshots', False):
                doc.add_heading('Snapshots', 1)
                
                if data['snapshots']:
                    # Create table
                    table = doc.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    
                    # Add table header
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'VM Name'
                    header_cells[1].text = 'Snapshot Name'
                    header_cells[2].text = 'Alter'
                    header_cells[3].text = 'Größe'
                    header_cells[4].text = 'Beschreibung'
                    
                    # Add table data
                    for snapshot in data['snapshots']:
                        row_cells = table.add_row().cells
                        row_cells[0].text = snapshot['vm_name']
                        row_cells[1].text = snapshot['name']
                        row_cells[2].text = snapshot.get('age_human', 'Unbekannt')
                        row_cells[3].text = str(snapshot.get('size_mb', 'Unbekannt')) + ' MB' if snapshot.get('size_mb') else 'Unbekannt'
                        row_cells[4].text = snapshot.get('description', '')
                        
                        # Highlight old snapshots
                        if snapshot.get('age_days', 0) > 30:
                            for paragraph in row_cells[2].paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor.from_string(self.colors['orange'][1:])
                else:
                    doc.add_paragraph("Keine Snapshot-Daten verfügbar.")
                
                doc.add_paragraph()
            
            # Add Orphaned VMDKs section if selected
            if options.get('orphaned_vmdks', False):
                doc.add_heading('Verwaiste VMDK-Dateien', 1)
                
                if data['orphaned_vmdks']:
                    # Create table
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    
                    # Add table header
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Dateiname'
                    header_cells[1].text = 'Datastore'
                    header_cells[2].text = 'Größe'
                    header_cells[3].text = 'Empfohlene Aktion'
                    
                    # Add table data
                    for vmdk in data['orphaned_vmdks']:
                        row_cells = table.add_row().cells
                        row_cells[0].text = vmdk['name']
                        row_cells[1].text = vmdk['datastore']
                        row_cells[2].text = vmdk.get('size_human', 'Unbekannt')
                        row_cells[3].text = vmdk.get('recommended_action', 'Manuelle Überprüfung erforderlich')
                else:
                    doc.add_paragraph("Keine verwaisten VMDK-Dateien gefunden.")
                
                doc.add_paragraph()
            
            # Add footer
            section = doc.sections[0]
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = f"Generiert mit VMware vSphere Reporter v29.0 | © 2025 Bechtle GmbH"
            
            # Save document
            doc.save(output_file)
            
            logger.info(f"Word-Bericht gespeichert: {output_file}")
        
        except Exception as e:
            logger.error(f"Fehler bei der Generierung des Word-Berichts: {str(e)}")
            raise