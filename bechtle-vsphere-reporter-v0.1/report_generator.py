"""
Bechtle vSphere Reporter - Report Generator
Generiert Berichte aus den gesammelten vSphere-Daten.
"""

import os
import logging
from datetime import datetime
import jinja2
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger('vsphere_reporter')

class ReportGenerator:
    """Generator für vSphere-Berichte in verschiedenen Formaten"""
    
    def __init__(self, data_collector):
        """
        Initialisiere den Report Generator
        
        Args:
            data_collector: DataCollector-Instanz mit Zugriff auf vSphere-Daten
        """
        self.data_collector = data_collector
        self.collector_available = data_collector is not None
        self.reports_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'reports')
        os.makedirs(self.reports_dir, exist_ok=True)
        
        # Bechtle-Farben
        self.colors = {
            'primary': '#00355e',    # Dunkelblau
            'secondary': '#da6f1e',  # Orange
            'success': '#23a96a',    # Grün
            'light': '#f3f3f3',      # Hellgrau
            'dark': '#5a5a5a'        # Dunkelgrau
        }
        
        # Jinja2-Template-Umgebung initialisieren
        template_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(template_dir),
            autoescape=jinja2.select_autoescape(['html', 'xml'])
        )
    
    def generate_report(self, options):
        """
        Generiert Berichte basierend auf den ausgewählten Optionen
        
        Args:
            options: Dictionary mit den ausgewählten Berichtsoptionen
            
        Returns:
            list: Liste der generierten Berichtsdateien
        """
        if not self.collector_available:
            logger.error("DataCollector nicht verfügbar")
            raise ValueError("DataCollector nicht verfügbar")
        
        # Timestamp für Dateinamen
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Sammle die erforderlichen Daten
        report_data = {}
        
        if options.get('vmware_tools', False):
            logger.info("Sammle VMware Tools-Daten")
            report_data['vmware_tools'] = self.data_collector.get_vmware_tools_status()
        
        if options.get('snapshots', False):
            logger.info("Sammle Snapshot-Daten")
            report_data['snapshots'] = self.data_collector.get_snapshot_info()
        
        if options.get('orphaned_vmdks', False):
            logger.info("Sammle Daten zu verwaisten VMDKs")
            report_data['orphaned_vmdks'] = self.data_collector.get_orphaned_vmdks()
        
        # Sammle Umgebungsdaten für Dashboard
        report_data['environment'] = self.data_collector.get_environment_stats()
        
        # Server-Informationen
        vs_client = self.data_collector.client
        report_data['server_info'] = {
            'name': vs_client.server if vs_client else 'Demo-Modus',
            'user': vs_client.username if vs_client else 'Demo-Benutzer',
            'timestamp': datetime.now().strftime('%d.%m.%Y %H:%M:%S'),
            'demo_mode': vs_client is None or not vs_client.is_connected()
        }
        
        # Generiere die ausgewählten Berichtsformate
        output_files = []
        
        if options.get('format_html', False):
            html_file = self._generate_html_report(report_data, timestamp)
            if html_file:
                output_files.append(html_file)
        
        if options.get('format_pdf', False):
            pdf_file = self._generate_pdf_report(report_data, timestamp)
            if pdf_file:
                output_files.append(pdf_file)
        
        if options.get('format_docx', False):
            docx_file = self._generate_docx_report(report_data, timestamp)
            if docx_file:
                output_files.append(docx_file)
        
        return output_files
    
    def _generate_html_report(self, report_data, timestamp):
        """
        Generiert einen HTML-Bericht
        
        Args:
            report_data: Die zu berichtenden Daten
            timestamp: Zeitstempel für den Dateinamen
            
        Returns:
            str: Pfad zur generierten Berichtsdatei oder None bei Fehler
        """
        try:
            template = self.jinja_env.get_template('report.html')
            
            # Render HTML mit Jinja2
            html_content = template.render(
                data=report_data,
                colors=self.colors,
                title='Bechtle vSphere Reporter - Bericht'
            )
            
            # Speichere HTML-Datei
            filename = f"vsphere_report_{timestamp}.html"
            file_path = os.path.join(self.reports_dir, filename)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
            logger.info(f"HTML-Bericht generiert: {file_path}")
            return filename
        except Exception as e:
            logger.error(f"Fehler bei der Generierung des HTML-Berichts: {str(e)}")
            return None
    
    def _generate_pdf_report(self, report_data, timestamp):
        """
        Generiert einen PDF-Bericht
        
        Args:
            report_data: Die zu berichtenden Daten
            timestamp: Zeitstempel für den Dateinamen
            
        Returns:
            str: Pfad zur generierten Berichtsdatei oder None bei Fehler
        """
        try:
            filename = f"vsphere_report_{timestamp}.pdf"
            file_path = os.path.join(self.reports_dir, filename)
            
            # Erstelle PDF mit ReportLab
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            elements = []
            
            # Titel
            title_style = styles['Title']
            title_style.textColor = colors.HexColor(self.colors['primary'])
            elements.append(Paragraph('Bechtle vSphere Reporter - Bericht', title_style))
            elements.append(Spacer(1, 12))
            
            # Server-Info
            server_info = report_data.get('server_info', {})
            server_text = f"Server: {server_info.get('name', 'N/A')}<br/>"
            server_text += f"Benutzer: {server_info.get('user', 'N/A')}<br/>"
            server_text += f"Datum: {server_info.get('timestamp', 'N/A')}"
            
            info_style = styles['Normal']
            elements.append(Paragraph(server_text, info_style))
            elements.append(Spacer(1, 24))
            
            # VMware Tools Status
            if 'vmware_tools' in report_data and report_data['vmware_tools']:
                heading_style = styles['Heading2']
                heading_style.textColor = colors.HexColor(self.colors['primary'])
                elements.append(Paragraph('VMware Tools Status', heading_style))
                elements.append(Spacer(1, 6))
                
                # Tabelle für VMware Tools
                data = [['VM-Name', 'Status', 'Version', 'Power State']]
                for tool in report_data['vmware_tools']:
                    data.append([
                        tool.get('vm_name', 'N/A'),
                        tool.get('tools_status', 'N/A'),
                        tool.get('tools_version', 'N/A'),
                        tool.get('power_state', 'N/A')
                    ])
                
                table = Table(data, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['primary'])),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                
                elements.append(table)
                elements.append(Spacer(1, 24))
            
            # Snapshots
            if 'snapshots' in report_data and report_data['snapshots']:
                elements.append(Paragraph('Snapshots', heading_style))
                elements.append(Spacer(1, 6))
                
                # Tabelle für Snapshots
                data = [['VM-Name', 'Snapshot-Name', 'Erstellungsdatum', 'Alter (Tage)']]
                for snap in report_data['snapshots']:
                    data.append([
                        snap.get('vm_name', 'N/A'),
                        snap.get('name', 'N/A'),
                        snap.get('create_time', 'N/A'),
                        str(snap.get('days_old', 'N/A'))
                    ])
                
                table = Table(data, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['primary'])),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                
                elements.append(table)
                elements.append(Spacer(1, 24))
            
            # Verwaiste VMDKs
            if 'orphaned_vmdks' in report_data and report_data['orphaned_vmdks']:
                elements.append(Paragraph('Verwaiste VMDK-Dateien', heading_style))
                elements.append(Spacer(1, 6))
                
                # Tabelle für VMDKs
                data = [['Dateiname', 'Datastore', 'Größe (GB)', 'Erstellungsdatum', 'Alter (Tage)']]
                for vmdk in report_data['orphaned_vmdks']:
                    data.append([
                        vmdk.get('name', 'N/A'),
                        vmdk.get('datastore', 'N/A'),
                        str(vmdk.get('size_gb', 'N/A')),
                        vmdk.get('creation_date', 'N/A'),
                        str(vmdk.get('days_old', 'N/A'))
                    ])
                
                table = Table(data, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.colors['primary'])),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                
                elements.append(table)
            
            # Footer
            elements.append(Spacer(1, 36))
            footer_text = f"Bericht generiert mit Bechtle vSphere Reporter v0.1 am {server_info.get('timestamp', 'N/A')}"
            footer_style = styles['Normal']
            footer_style.textColor = colors.HexColor(self.colors['dark'])
            footer_style.alignment = 1  # Zentriert
            elements.append(Paragraph(footer_text, footer_style))
            
            # Build PDF
            doc.build(elements)
            
            logger.info(f"PDF-Bericht generiert: {file_path}")
            return filename
        except Exception as e:
            logger.error(f"Fehler bei der Generierung des PDF-Berichts: {str(e)}")
            return None
    
    def _generate_docx_report(self, report_data, timestamp):
        """
        Generiert einen DOCX-Bericht
        
        Args:
            report_data: Die zu berichtenden Daten
            timestamp: Zeitstempel für den Dateinamen
            
        Returns:
            str: Pfad zur generierten Berichtsdatei oder None bei Fehler
        """
        try:
            filename = f"vsphere_report_{timestamp}.docx"
            file_path = os.path.join(self.reports_dir, filename)
            
            # Erstelle DOCX mit python-docx
            doc = Document()
            
            # Dokumenteigenschaften
            doc.core_properties.title = "Bechtle vSphere Reporter - Bericht"
            doc.core_properties.author = "Bechtle vSphere Reporter"
            
            # Seitenränder einstellen
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
            
            # Titel
            title = doc.add_heading('Bechtle vSphere Reporter - Bericht', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Server-Info
            server_info = report_data.get('server_info', {})
            server_para = doc.add_paragraph()
            server_para.add_run(f"Server: {server_info.get('name', 'N/A')}\n")
            server_para.add_run(f"Benutzer: {server_info.get('user', 'N/A')}\n")
            server_para.add_run(f"Datum: {server_info.get('timestamp', 'N/A')}")
            doc.add_paragraph()
            
            # VMware Tools Status
            if 'vmware_tools' in report_data and report_data['vmware_tools']:
                doc.add_heading('VMware Tools Status', level=1)
                
                # Tabelle für VMware Tools
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Header-Zeile
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'VM-Name'
                hdr_cells[1].text = 'Status'
                hdr_cells[2].text = 'Version'
                hdr_cells[3].text = 'Power State'
                
                # Formatierung der Header-Zeile
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(12)
                
                # Daten-Zeilen
                for tool in report_data['vmware_tools']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = tool.get('vm_name', 'N/A')
                    row_cells[1].text = tool.get('tools_status', 'N/A')
                    row_cells[2].text = tool.get('tools_version', 'N/A')
                    row_cells[3].text = tool.get('power_state', 'N/A')
                
                doc.add_paragraph()
            
            # Snapshots
            if 'snapshots' in report_data and report_data['snapshots']:
                doc.add_heading('Snapshots', level=1)
                
                # Tabelle für Snapshots
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Header-Zeile
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'VM-Name'
                hdr_cells[1].text = 'Snapshot-Name'
                hdr_cells[2].text = 'Erstellungsdatum'
                hdr_cells[3].text = 'Alter (Tage)'
                
                # Formatierung der Header-Zeile
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(12)
                
                # Daten-Zeilen
                for snap in report_data['snapshots']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = snap.get('vm_name', 'N/A')
                    row_cells[1].text = snap.get('name', 'N/A')
                    row_cells[2].text = snap.get('create_time', 'N/A')
                    row_cells[3].text = str(snap.get('days_old', 'N/A'))
                
                doc.add_paragraph()
            
            # Verwaiste VMDKs
            if 'orphaned_vmdks' in report_data and report_data['orphaned_vmdks']:
                doc.add_heading('Verwaiste VMDK-Dateien', level=1)
                
                # Tabelle für VMDKs
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Header-Zeile
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Dateiname'
                hdr_cells[1].text = 'Datastore'
                hdr_cells[2].text = 'Größe (GB)'
                hdr_cells[3].text = 'Erstellungsdatum'
                hdr_cells[4].text = 'Alter (Tage)'
                
                # Formatierung der Header-Zeile
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(12)
                
                # Daten-Zeilen
                for vmdk in report_data['orphaned_vmdks']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = vmdk.get('name', 'N/A')
                    row_cells[1].text = vmdk.get('datastore', 'N/A')
                    row_cells[2].text = str(vmdk.get('size_gb', 'N/A'))
                    row_cells[3].text = vmdk.get('creation_date', 'N/A')
                    row_cells[4].text = str(vmdk.get('days_old', 'N/A'))
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph(f"Bericht generiert mit Bechtle vSphere Reporter v0.1 am {server_info.get('timestamp', 'N/A')}")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save DOCX
            doc.save(file_path)
            
            logger.info(f"DOCX-Bericht generiert: {file_path}")
            return filename
        except Exception as e:
            logger.error(f"Fehler bei der Generierung des DOCX-Berichts: {str(e)}")
            return None