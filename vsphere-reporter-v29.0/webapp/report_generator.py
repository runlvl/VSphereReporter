#!/usr/bin/env python3
"""
Report Generator Module for VMware vSphere Reporter Web Edition
Generates comprehensive reports in various formats
"""

import os
import logging
import traceback
import humanize
from datetime import datetime
from jinja2 import Environment, FileSystemLoader
from webapp.direct_vmdk_collector import DirectVMDKCollector
from webapp.topology_generator import TopologyGenerator

# Configure logger
logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generator for vSphere environment reports"""
    
    def __init__(self, data_collector):
        """
        Initialize the report generator
        
        Args:
            data_collector: Configured DataCollector instance
        """
        self.collector = data_collector
        self.client = data_collector.client
        self.logger = logging.getLogger(__name__)
        self.template_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'templates')
    
    def generate_html_report(self, output_file, include_vmware_tools=True, include_snapshots=True,
                             include_orphaned_vmdks=True, include_vm_hardware=True, include_datastores=True,
                             include_clusters=True, include_hosts=True, include_resource_pools=True,
                             include_networks=True, include_topology=True):
        """
        Generate an HTML report
        
        Args:
            output_file: Path to output file
            include_*: Flags to control which sections to include
            
        Returns:
            str: Path to generated file
        """
        if not self.client.is_connected():
            raise Exception("Not connected to vCenter Server")
        
        self.logger.info(f"Generating HTML report: {output_file}")
        
        try:
            # Collect required data based on selected options
            report_data = self._collect_report_data(
                include_vmware_tools, include_snapshots, include_orphaned_vmdks,
                include_vm_hardware, include_datastores, include_clusters,
                include_hosts, include_resource_pools, include_networks, include_topology
            )
            
            # Load template environment
            env = Environment(loader=FileSystemLoader(self.template_dir))
            template = env.get_template('report_template.html')
            
            # Render template with data
            html_content = template.render(**report_data)
            
            # Write HTML to file
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.logger.info(f"HTML report generated successfully: {output_file}")
            return output_file
        
        except Exception as e:
            self.logger.error(f"Failed to generate HTML report: {str(e)}")
            self.logger.debug(f"Traceback: {traceback.format_exc()}")
            
            # Try to generate a simple fallback report
            return self._generate_simple_html_report(
                output_file, str(e), include_vmware_tools, include_snapshots, include_orphaned_vmdks
            )
    
    def generate_pdf_report(self, output_file, include_vmware_tools=True, include_snapshots=True,
                           include_orphaned_vmdks=True, include_vm_hardware=True, include_datastores=True,
                           include_clusters=True, include_hosts=True, include_resource_pools=True,
                           include_networks=True, include_topology=True):
        """
        Generate a PDF report
        
        Args:
            output_file: Path to output file
            include_*: Flags to control which sections to include
            
        Returns:
            str: Path to generated file
        """
        if not self.client.is_connected():
            raise Exception("Not connected to vCenter Server")
        
        self.logger.info(f"Generating PDF report: {output_file}")
        
        try:
            # First, generate HTML report to a temporary file
            temp_html_file = output_file.replace('.pdf', '_temp.html')
            self.generate_html_report(
                temp_html_file, include_vmware_tools, include_snapshots, include_orphaned_vmdks,
                include_vm_hardware, include_datastores, include_clusters, include_hosts,
                include_resource_pools, include_networks, include_topology
            )
            
            # Convert HTML to PDF using a dedicated HTML to PDF converter
            # This requires an external library like weasyprint or wkhtmltopdf
            # For this example, we'll use a placeholder implementation
            self._html_to_pdf(temp_html_file, output_file)
            
            # Clean up temporary file
            if os.path.exists(temp_html_file):
                os.remove(temp_html_file)
            
            self.logger.info(f"PDF report generated successfully: {output_file}")
            return output_file
        
        except Exception as e:
            self.logger.error(f"Failed to generate PDF report: {str(e)}")
            self.logger.debug(f"Traceback: {traceback.format_exc()}")
            
            # Fallback to HTML report if PDF generation fails
            html_fallback = output_file.replace('.pdf', '_fallback.html')
            self.logger.warning(f"Falling back to HTML report: {html_fallback}")
            return self._generate_simple_html_report(
                html_fallback, f"PDF generation failed: {str(e)}",
                include_vmware_tools, include_snapshots, include_orphaned_vmdks
            )
    
    def generate_docx_report(self, output_file, include_vmware_tools=True, include_snapshots=True,
                            include_orphaned_vmdks=True, include_vm_hardware=True, include_datastores=True,
                            include_clusters=True, include_hosts=True, include_resource_pools=True,
                            include_networks=True, include_topology=True):
        """
        Generate a DOCX report
        
        Args:
            output_file: Path to output file
            include_*: Flags to control which sections to include
            
        Returns:
            str: Path to generated file
        """
        if not self.client.is_connected():
            raise Exception("Not connected to vCenter Server")
        
        self.logger.info(f"Generating DOCX report: {output_file}")
        
        try:
            # Collect required data based on selected options
            report_data = self._collect_report_data(
                include_vmware_tools, include_snapshots, include_orphaned_vmdks,
                include_vm_hardware, include_datastores, include_clusters,
                include_hosts, include_resource_pools, include_networks, include_topology
            )
            
            # Generate DOCX using python-docx
            # This requires the python-docx library
            # For this example, we'll use a placeholder implementation
            self._generate_docx(output_file, report_data)
            
            self.logger.info(f"DOCX report generated successfully: {output_file}")
            return output_file
        
        except Exception as e:
            self.logger.error(f"Failed to generate DOCX report: {str(e)}")
            self.logger.debug(f"Traceback: {traceback.format_exc()}")
            
            # Fallback to HTML report if DOCX generation fails
            html_fallback = output_file.replace('.docx', '_fallback.html')
            self.logger.warning(f"Falling back to HTML report: {html_fallback}")
            return self._generate_simple_html_report(
                html_fallback, f"DOCX generation failed: {str(e)}",
                include_vmware_tools, include_snapshots, include_orphaned_vmdks
            )
    
    def _collect_report_data(self, include_vmware_tools, include_snapshots, include_orphaned_vmdks,
                            include_vm_hardware, include_datastores, include_clusters,
                            include_hosts, include_resource_pools, include_networks, include_topology):
        """
        Collect data for the report based on selected options
        
        Args:
            include_*: Flags to control which data to collect
            
        Returns:
            dict: Data for the report
        """
        report_data = {
            'title': f"VMware vSphere Infrastructure Report",
            'subtitle': f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            'vcenter_server': self.client.host,
            'vcenter_version': self.client.content.about.version,
            'vcenter_build': self.client.content.about.build,
            'vcenter_os': self.client.content.about.osType,
        }
        
        # Collect VMware Tools information
        if include_vmware_tools:
            self.logger.info("Collecting VMware Tools information")
            report_data['vmware_tools'] = self.collector.collect_vmware_tools_info()
        
        # Collect snapshot information
        if include_snapshots:
            self.logger.info("Collecting snapshot information")
            report_data['snapshots'] = self.collector.collect_snapshot_info()
        
        # Collect VMDK information
        if include_orphaned_vmdks:
            self.logger.info("Collecting VMDK information")
            vmdk_collector = DirectVMDKCollector(self.client)
            report_data['vmdks'] = vmdk_collector.collect_all_vmdks()
            report_data['orphaned_vmdks'] = [vmdk for vmdk in report_data['vmdks'] if vmdk['status'] == 'POTENTIALLY ORPHANED']
        
        # Collect VM hardware information
        if include_vm_hardware:
            self.logger.info("Collecting VM hardware information")
            report_data['vm_hardware'] = self.collector.collect_vm_hardware_info()
        
        # Collect datastore information
        if include_datastores:
            self.logger.info("Collecting datastore information")
            report_data['datastores'] = self.collector.collect_datastore_info()
        
        # Collect cluster information
        if include_clusters:
            self.logger.info("Collecting cluster information")
            report_data['clusters'] = self.collector.collect_cluster_info()
        
        # Collect host information
        if include_hosts:
            self.logger.info("Collecting host information")
            report_data['hosts'] = self.collector.collect_host_info()
        
        # Collect resource pool information
        if include_resource_pools:
            self.logger.info("Collecting resource pool information")
            report_data['resource_pools'] = self.collector.collect_resource_pool_info()
        
        # Collect network information
        if include_networks:
            self.logger.info("Collecting network information")
            report_data['networks'] = self.collector.collect_network_info()
        
        # Generate topology chart
        if include_topology:
            self.logger.info("Generating topology chart")
            topology_generator = TopologyGenerator(self.client)
            report_data['topology_chart'] = topology_generator.generate_topology_chart()
        
        return report_data
    
    def _generate_simple_html_report(self, output_file, error_message, include_vmware_tools, include_snapshots, include_orphaned_vmdks):
        """
        Generate a simple HTML report as a fallback
        
        Args:
            output_file: Path to output file
            error_message: Error message to include in the report
            include_*: Flags to control which sections to include
            
        Returns:
            str: Path to generated file
        """
        self.logger.info(f"Generating simple fallback HTML report: {output_file}")
        
        try:
            html_parts = []
            html_parts.append('<!DOCTYPE html>')
            html_parts.append('<html lang="de">')
            html_parts.append('<head>')
            html_parts.append('    <meta charset="UTF-8">')
            html_parts.append('    <meta name="viewport" content="width=device-width, initial-scale=1.0">')
            html_parts.append('    <title>VMware vSphere Report (Fallback Version)</title>')
            html_parts.append('    <style>')
            html_parts.append('        body { font-family: Arial, sans-serif; line-height: 1.6; color: #333; max-width: 1200px; margin: 0 auto; padding: 20px; }')
            html_parts.append('        h1, h2, h3 { color: #00355e; }')
            html_parts.append('        h1 { border-bottom: 2px solid #00355e; padding-bottom: 10px; }')
            html_parts.append('        h2 { border-bottom: 1px solid #da6f1e; padding-bottom: 5px; margin-top: 30px; }')
            html_parts.append('        table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }')
            html_parts.append('        th, td { text-align: left; padding: 8px; border: 1px solid #ddd; }')
            html_parts.append('        th { background-color: #f3f3f3; }')
            html_parts.append('        tr:nth-child(even) { background-color: #f9f9f9; }')
            html_parts.append('        .error { color: #cc0000; background-color: #ffeeee; padding: 10px; border-left: 5px solid #cc0000; margin-bottom: 20px; }')
            html_parts.append('        .warning { color: #ff6600; background-color: #fff8ee; padding: 10px; border-left: 5px solid #ff6600; margin-bottom: 20px; }')
            html_parts.append('        .info { color: #006699; background-color: #eef8ff; padding: 10px; border-left: 5px solid #006699; margin-bottom: 20px; }')
            html_parts.append('        .status-potentially-orphaned { color: #cc0000; font-weight: bold; }')
            html_parts.append('        .status-aktiv { color: #009900; }')
            html_parts.append('        .status-template { color: #0066cc; }')
            html_parts.append('        .status-helper { color: #888888; font-style: italic; }')
            html_parts.append('    </style>')
            html_parts.append('</head>')
            html_parts.append('<body>')
            
            # Header
            html_parts.append('    <h1>VMware vSphere Infrastructure Report (Fallback Version)</h1>')
            html_parts.append(f'    <p>Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>')
            html_parts.append(f'    <p>vCenter Server: {self.client.host}</p>')
            
            # Error message
            html_parts.append('    <div class="error">')
            html_parts.append('        <h3>Error during report generation</h3>')
            html_parts.append(f'        <p>{error_message}</p>')
            html_parts.append('        <p>This is a simplified fallback report. Some information may be missing or incomplete.</p>')
            html_parts.append('    </div>')
            
            # Content based on available data
            if include_vmware_tools:
                try:
                    vmware_tools = self.collector.collect_vmware_tools_info()
                    if vmware_tools:
                        html_parts.append('    <h2>VMware Tools Status</h2>')
                        html_parts.append('    <table>')
                        html_parts.append('        <tr>')
                        html_parts.append('            <th>VM Name</th>')
                        html_parts.append('            <th>Power State</th>')
                        html_parts.append('            <th>Tools Version</th>')
                        html_parts.append('            <th>Tools Status</th>')
                        html_parts.append('            <th>Running Status</th>')
                        html_parts.append('        </tr>')
                        for vm in vmware_tools:
                            html_parts.append('        <tr>')
                            html_parts.append(f'            <td>{vm["name"]}</td>')
                            html_parts.append(f'            <td>{vm["power_state"]}</td>')
                            html_parts.append(f'            <td>{vm["tools_version"]}</td>')
                            html_parts.append(f'            <td>{vm["tools_status"]}</td>')
                            html_parts.append(f'            <td>{vm["tools_running_status"]}</td>')
                            html_parts.append('        </tr>')
                        html_parts.append('    </table>')
                except Exception as e:
                    self.logger.warning(f"Failed to include VMware Tools info in fallback report: {str(e)}")
                    html_parts.append('    <h2>VMware Tools Status</h2>')
                    html_parts.append(f'    <div class="warning">Failed to collect VMware Tools information: {str(e)}</div>')
            
            if include_snapshots:
                try:
                    snapshots = self.collector.collect_snapshot_info()
                    if snapshots:
                        html_parts.append('    <h2>VM Snapshots</h2>')
                        html_parts.append('    <table>')
                        html_parts.append('        <tr>')
                        html_parts.append('            <th>VM Name</th>')
                        html_parts.append('            <th>Snapshot Name</th>')
                        html_parts.append('            <th>Description</th>')
                        html_parts.append('            <th>Age</th>')
                        html_parts.append('            <th>Created</th>')
                        html_parts.append('            <th>Size (MB)</th>')
                        html_parts.append('        </tr>')
                        for snapshot in snapshots:
                            html_parts.append('        <tr>')
                            html_parts.append(f'            <td>{snapshot["vm_name"]}</td>')
                            html_parts.append(f'            <td>{snapshot["name"]}</td>')
                            html_parts.append(f'            <td>{snapshot["description"]}</td>')
                            html_parts.append(f'            <td>{snapshot["age_text"]}</td>')
                            html_parts.append(f'            <td>{snapshot["create_time"].strftime("%Y-%m-%d %H:%M:%S")}</td>')
                            html_parts.append(f'            <td>{snapshot["size_mb"]:.2f}</td>')
                            html_parts.append('        </tr>')
                        html_parts.append('    </table>')
                except Exception as e:
                    self.logger.warning(f"Failed to include snapshot info in fallback report: {str(e)}")
                    html_parts.append('    <h2>VM Snapshots</h2>')
                    html_parts.append(f'    <div class="warning">Failed to collect snapshot information: {str(e)}</div>')
            
            if include_orphaned_vmdks:
                try:
                    vmdk_collector = DirectVMDKCollector(self.client)
                    vmdks = vmdk_collector.collect_all_vmdks()
                    if vmdks:
                        html_parts.append('    <h2>VMDK Status</h2>')
                        html_parts.append('    <table>')
                        html_parts.append('        <tr>')
                        html_parts.append('            <th>Datastore</th>')
                        html_parts.append('            <th>VMDK Path</th>')
                        html_parts.append('            <th>Size</th>')
                        html_parts.append('            <th>Status</th>')
                        html_parts.append('            <th>VM Name</th>')
                        html_parts.append('        </tr>')
                        for vmdk in vmdks:
                            status_class = ''
                            if vmdk['status'] == 'POTENTIALLY ORPHANED':
                                status_class = 'status-potentially-orphaned'
                            elif vmdk['status'] == 'AKTIV':
                                status_class = 'status-aktiv'
                            elif vmdk['status'] == 'TEMPLATE':
                                status_class = 'status-template'
                            elif vmdk['status'] == 'HELPER':
                                status_class = 'status-helper'
                            
                            html_parts.append('        <tr>')
                            html_parts.append(f'            <td>{vmdk["datastore"]}</td>')
                            html_parts.append(f'            <td>{vmdk["path"]}</td>')
                            html_parts.append(f'            <td>{humanize.naturalsize(vmdk["size_kb"] * 1024)}</td>')
                            html_parts.append(f'            <td class="{status_class}">{vmdk["status"]}</td>')
                            html_parts.append(f'            <td>{vmdk["vm_name"] or "N/A"}</td>')
                            html_parts.append('        </tr>')
                        html_parts.append('    </table>')
                except Exception as e:
                    self.logger.warning(f"Failed to include VMDK info in fallback report: {str(e)}")
                    html_parts.append('    <h2>VMDK Status</h2>')
                    html_parts.append(f'    <div class="warning">Failed to collect VMDK information: {str(e)}</div>')
            
            # Footer
            html_parts.append('    <div class="info">')
            html_parts.append('        <p><strong>Note:</strong> This is a simplified fallback report generated due to an error in the standard report generation process.</p>')
            html_parts.append('        <p>Please try generating the report again with fewer options selected if you encounter consistent issues.</p>')
            html_parts.append('    </div>')
            
            html_parts.append('</body>')
            html_parts.append('</html>')
            
            # Write to file
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html_parts))
            
            self.logger.info(f"Simple fallback HTML report generated successfully: {output_file}")
            return output_file
        
        except Exception as e:
            self.logger.error(f"Failed to generate simple fallback HTML report: {str(e)}")
            self.logger.debug(f"Traceback: {traceback.format_exc()}")
            
            # Last resort - create an ultra-simple error report
            error_file = output_file.replace('.html', '_error.html')
            with open(error_file, 'w', encoding='utf-8') as f:
                f.write(f'<!DOCTYPE html><html><head><title>Error Report</title></head><body>')
                f.write(f'<h1>Critical Error in Report Generation</h1>')
                f.write(f'<p>Error message: {str(e)}</p>')
                f.write(f'<p>Original error: {error_message}</p>')
                f.write(f'<p>Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>')
                f.write(f'</body></html>')
            
            return error_file
    
    def _html_to_pdf(self, html_file, pdf_file):
        """
        Convert HTML to PDF
        
        Args:
            html_file: Path to HTML file
            pdf_file: Path to output PDF file
        """
        # This is a placeholder implementation
        # In a real implementation, this would use a library like weasyprint or wkhtmltopdf
        # For now, we'll just copy the HTML file to a .pdf extension as a demonstration
        self.logger.warning("HTML to PDF conversion is not fully implemented in this version")
        self.logger.warning("In a production environment, this would use a proper HTML to PDF converter")
        
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        c = canvas.Canvas(pdf_file, pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(30, 750, "VMware vSphere Infrastructure Report")
        c.drawString(30, 730, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        c.drawString(30, 700, f"vCenter Server: {self.client.host}")
        c.drawString(30, 680, "PDF generation is not fully implemented in this version.")
        c.drawString(30, 660, "Please use the HTML report for complete information.")
        c.save()
    
    def _generate_docx(self, docx_file, report_data):
        """
        Generate a DOCX report
        
        Args:
            docx_file: Path to output DOCX file
            report_data: Data for the report
        """
        # This is a placeholder implementation
        # In a real implementation, this would use the python-docx library
        self.logger.warning("DOCX generation is not fully implemented in this version")
        self.logger.warning("In a production environment, this would use the python-docx library")
        
        import docx
        from docx.shared import Pt, RGBColor, Inches
        
        doc = docx.Document()
        
        # Add heading
        doc.add_heading("VMware vSphere Infrastructure Report", 0)
        doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"vCenter Server: {report_data['vcenter_server']}")
        doc.add_paragraph(f"vCenter Version: {report_data['vcenter_version']} (Build {report_data['vcenter_build']})")
        
        # Add a note about this being a placeholder
        p = doc.add_paragraph("Note: ")
        p.add_run("DOCX report generation is not fully implemented in this version. Please use the HTML report for complete information.").italic = True
        
        # Add basic sections
        if 'vmware_tools' in report_data and report_data['vmware_tools']:
            doc.add_heading("VMware Tools Status", 1)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'VM Name'
            hdr_cells[1].text = 'Power State'
            hdr_cells[2].text = 'Tools Version'
            hdr_cells[3].text = 'Tools Status'
            
            for vm in report_data['vmware_tools'][:10]:  # Just show first 10 for placeholder
                row_cells = table.add_row().cells
                row_cells[0].text = vm['name']
                row_cells[1].text = vm['power_state']
                row_cells[2].text = vm['tools_version']
                row_cells[3].text = vm['tools_status']
            
            doc.add_paragraph("(Showing first 10 VMs only in this placeholder implementation)")
        
        # Save the document
        doc.save(docx_file)