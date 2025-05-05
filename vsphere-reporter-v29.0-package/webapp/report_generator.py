#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
VMware vSphere Reporter v29.0 - Report Generator

Dieses Modul stellt Funktionen für die Generierung von Berichten
in verschiedenen Formaten wie HTML, PDF und DOCX bereit.

Copyright (c) 2025 Bechtle GmbH
"""

import os
import sys
import logging
import json
import re
from datetime import datetime
import tempfile
import shutil

# Versuche, die Berichtsmodule zu importieren
try:
    # Für PDF-Berichte
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Für DOCX-Berichte
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    # Logge Warnungen, falls die Berichtsmodule nicht verfügbar sind
    logging.warning("Einige Berichtsmodule konnten nicht importiert werden. "
                   "PDF- oder DOCX-Berichte sind möglicherweise nicht verfügbar.")

# Lokale Module importieren
from webapp.utils.error_handler import handle_vsphere_errors, VSphereReportGenerationError

logger = logging.getLogger(__name__)

# Bechtle Farben
BECHTLE_DARK_BLUE = "#00355e"
BECHTLE_ORANGE = "#da6f1e"
BECHTLE_GREEN = "#23a96a"
BECHTLE_LIGHT_GRAY = "#f3f3f3"
BECHTLE_DARK_GRAY = "#5a5a5a"

class ReportGenerator:
    """
    Generator für VMware vSphere-Berichte in verschiedenen Formaten.
    
    Unterstützt die Generierung von Berichten im HTML-, PDF- und DOCX-Format
    mit Informationen über VMware Tools, Snapshots und verwaiste VMDKs.
    """
    
    def __init__(self, output_dir, vcenter_server=None, username=None, demo_mode=False):
        """
        Initialisiert den Report-Generator.
        
        Args:
            output_dir (str): Verzeichnis für die Ausgabedateien
            vcenter_server (str, optional): Name des vCenter-Servers
            username (str, optional): Benutzername, der die Berichte erstellt
            demo_mode (bool, optional): Gibt an, ob der Demo-Modus aktiv ist
        """
        self.output_dir = output_dir
        self.vcenter_server = vcenter_server or 'Unknown vCenter'
        self.username = username or 'Unknown User'
        self.demo_mode = demo_mode
        
        # Stelle sicher, dass das Ausgabeverzeichnis existiert
        os.makedirs(output_dir, exist_ok=True)
        
        # Zeitstempel für den Bericht
        self.timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.datestamp = datetime.now().strftime("%Y-%m-%d")
    
    @handle_vsphere_errors
    def generate_html_report(self, filename, report_data):
        """
        Generiert einen HTML-Bericht.
        
        Args:
            filename (str): Name der Ausgabedatei
            report_data (dict): Daten für den Bericht:
                - vmware_tools: Liste von VMware Tools-Informationen
                - snapshots: Liste von Snapshot-Informationen
                - orphaned_vmdks: Liste von verwaisten VMDK-Informationen
                
        Returns:
            str: Pfad zur generierten HTML-Datei
        """
        logger.info(f"Generiere HTML-Bericht: {filename}")
        
        try:
            # Verzeichnispfad erstellen
            file_path = os.path.join(self.output_dir, filename)
            
            # Einfaches HTML-Template mit Bootstrap erstellen
            with open(file_path, 'w', encoding='utf-8') as f:
                # HTML-Header
                f.write(f"""<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>VMware vSphere Bericht - {self.datestamp}</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0-beta3/css/all.min.css">
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; }}
        .header {{ background-color: #00355e; color: white; padding: 20px; }}
        .report-section {{ margin-bottom: 30px; }}
        .footer {{ background-color: #f8f9fa; padding: 10px; font-size: 0.8em; }}
        .status-ok {{ color: #198754; }}
        .status-warning {{ color: #ffc107; }}
        .status-error {{ color: #dc3545; }}
        .table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
        .table th {{ background-color: #f3f3f3; }}
    </style>
</head>
<body>
    <div class="header">
        <div class="container">
            <div class="row align-items-center">
                <div class="col-md-9">
                    <h1>VMware vSphere Umgebungsbericht</h1>
                    <p>Erstellt am {self.timestamp}</p>
                </div>
                <div class="col-md-3 text-end">
                    <img src="https://www.bechtle.com/dam/jcr:6edbcbdc-6dc2-4c6f-8d9f-f7e3606d466a/bechtle_logo_weiss.svg" alt="Bechtle Logo" height="60">
                </div>
            </div>
        </div>
    </div>
    
    <div class="container my-4">
        <div class="alert alert-info mb-4">
            <strong>vCenter Server:</strong> {self.vcenter_server} | 
            <strong>Benutzer:</strong> {self.username} |
            <strong>Berichtdatum:</strong> {self.datestamp}
            {' <span class="badge bg-warning">Demo-Modus</span>' if self.demo_mode else ''}
        </div>
""")
                
                # VMware Tools Bericht
                if 'vmware_tools' in report_data and report_data['vmware_tools']:
                    vmware_tools = report_data['vmware_tools']
                    
                    f.write("""
        <div class="report-section">
            <h2><i class="fas fa-tools me-2"></i>VMware Tools Status</h2>
            <div class="card shadow-sm">
                <div class="card-body">
                    <p>VMware Tools-Status für alle virtuellen Maschinen. 
                    VMs mit veralteten Tools sind möglicherweise nicht kompatibel mit neuen Funktionen.</p>
                    
                    <table class="table table-striped table-hover">
                        <thead>
                            <tr>
                                <th>VM-Name</th>
                                <th>Betriebssystem</th>
                                <th>Status</th>
                                <th>Version</th>
                                <th>Laufzeitstatus</th>
                                <th>Letzte Aktualisierung</th>
                                <th>VM-Power</th>
                            </tr>
                        </thead>
                        <tbody>
""")
                    
                    # Sortieren nach Status (Problem first)
                    sorted_vms = sorted(vmware_tools, key=lambda x: (
                        0 if x.get('tools_status') == 'NotInstalled' else
                        1 if x.get('tools_status') == 'UpdateNeeded' else
                        2 if x.get('tools_status') == 'NotRunning' else
                        3 if x.get('tools_status') == 'Unmanaged' else 4
                    ))
                    
                    for vm in sorted_vms:
                        status_class = ""
                        status_text = ""
                        
                        if vm.get('tools_status') == 'Current':
                            status_class = "status-ok"
                            status_text = "Aktuell"
                        elif vm.get('tools_status') == 'UpdateNeeded':
                            status_class = "status-warning"
                            status_text = "Update verfügbar"
                        elif vm.get('tools_status') == 'NotInstalled':
                            status_class = "status-error"
                            status_text = "Nicht installiert"
                        elif vm.get('tools_status') == 'NotRunning':
                            status_class = "status-warning"
                            status_text = "Nicht ausgeführt"
                        else:
                            status_text = "Unbekannt"
                        
                        f.write(f"""
                            <tr>
                                <td>{vm.get('vm_name', 'Unbekannt')}</td>
                                <td>{vm.get('os', 'Unbekannt')}</td>
                                <td class="{status_class}">{status_text}</td>
                                <td>{vm.get('tools_version', 'Unbekannt')}</td>
                                <td>{vm.get('tools_running_status', 'Unbekannt')}</td>
                                <td>{vm.get('tools_update_time', 'Unbekannt')}</td>
                                <td>{'Eingeschaltet' if vm.get('power_state') == 'poweredOn' else 'Ausgeschaltet'}</td>
                            </tr>
                        """)
                    
                    f.write("""
                        </tbody>
                    </table>
                </div>
            </div>
        </div>
""")
                
                # Snapshots Bericht
                if 'snapshots' in report_data and report_data['snapshots']:
                    snapshots = report_data['snapshots']
                    
                    f.write("""
        <div class="report-section">
            <h2><i class="fas fa-camera me-2"></i>VM-Snapshots</h2>
            <div class="card shadow-sm">
                <div class="card-body">
                    <p>Liste aller aktiven Snapshots in der Umgebung. 
                    Ältere Snapshots könnten zu Leistungsproblemen führen oder Speicherplatz vergeuden.</p>
                    
                    <table class="table table-striped table-hover">
                        <thead>
                            <tr>
                                <th>VM-Name</th>
                                <th>Snapshot-Name</th>
                                <th>Beschreibung</th>
                                <th>Erstellt am</th>
                                <th>Alter (Tage)</th>
                                <th>Größe</th>
                            </tr>
                        </thead>
                        <tbody>
""")
                    
                    # Sortieren nach Alter (älteste zuerst)
                    sorted_snapshots = sorted(snapshots, key=lambda x: x.get('days_old', 0), reverse=True)
                    
                    for snapshot in sorted_snapshots:
                        age_class = ""
                        if snapshot.get('days_old', 0) > 30:
                            age_class = "status-error"
                        elif snapshot.get('days_old', 0) > 14:
                            age_class = "status-warning"
                        
                        f.write(f"""
                            <tr>
                                <td>{snapshot.get('vm_name', 'Unbekannt')}</td>
                                <td>{snapshot.get('name', 'Unbekannt')}</td>
                                <td>{snapshot.get('description', '')}</td>
                                <td>{snapshot.get('created_date', 'Unbekannt')}</td>
                                <td class="{age_class}">{snapshot.get('days_old', 'Unbekannt')}</td>
                                <td>{snapshot.get('size_gb', 'Unbekannt')} GB</td>
                            </tr>
                        """)
                    
                    f.write("""
                        </tbody>
                    </table>
                </div>
            </div>
        </div>
""")
                
                # Verwaiste VMDKs Bericht
                if 'orphaned_vmdks' in report_data and report_data['orphaned_vmdks']:
                    orphaned_vmdks = report_data['orphaned_vmdks']
                    
                    f.write("""
        <div class="report-section">
            <h2><i class="fas fa-hdd me-2"></i>Verwaiste VMDK-Dateien</h2>
            <div class="card shadow-sm">
                <div class="card-body">
                    <p>Liste aller verwaisten VMDK-Dateien in der Umgebung.
                    Diese belegen Speicherplatz, werden aber von keiner VM verwendet.</p>
                    
                    <table class="table table-striped table-hover">
                        <thead>
                            <tr>
                                <th>Datastore</th>
                                <th>Pfad</th>
                                <th>Größe (GB)</th>
                                <th>Letzte Änderung</th>
                                <th>Alter (Tage)</th>
                                <th>Vermutete VM</th>
                                <th>Empfohlene Aktion</th>
                            </tr>
                        </thead>
                        <tbody>
""")
                    
                    # Sortieren nach Größe (größte zuerst)
                    sorted_vmdks = sorted(orphaned_vmdks, key=lambda x: x.get('size_gb', 0), reverse=True)
                    
                    for vmdk in sorted_vmdks:
                        action_class = ""
                        if vmdk.get('recovery_action') == "Kann gelöscht werden":
                            action_class = "status-error"
                        elif vmdk.get('recovery_action') == "Überprüfung empfohlen":
                            action_class = "status-warning"
                        
                        f.write(f"""
                            <tr>
                                <td>{vmdk.get('datastore', 'Unbekannt')}</td>
                                <td>{vmdk.get('path', 'Unbekannt')}</td>
                                <td>{vmdk.get('size_gb', 'Unbekannt')} GB</td>
                                <td>{vmdk.get('last_modified', 'Unbekannt')}</td>
                                <td>{vmdk.get('days_orphaned', 'Unbekannt')}</td>
                                <td>{vmdk.get('probable_vm', 'Unbekannt')}</td>
                                <td class="{action_class}">{vmdk.get('recovery_action', 'Unbekannt')}</td>
                            </tr>
                        """)
                    
                    f.write("""
                        </tbody>
                    </table>
                </div>
            </div>
        </div>
""")
                
                # Zusammenfassung
                f.write("""
        <div class="alert alert-primary" role="alert">
            <h4><i class="fas fa-clipboard-check me-2"></i>Zusammenfassung</h4>
            <ul>
""")
                
                if 'vmware_tools' in report_data:
                    outdated_tools = sum(1 for vm in report_data['vmware_tools'] if vm.get('tools_status') == 'UpdateNeeded')
                    missing_tools = sum(1 for vm in report_data['vmware_tools'] if vm.get('tools_status') == 'NotInstalled')
                    f.write(f"<li>{outdated_tools} VMs mit veralteten VMware Tools</li>")
                    f.write(f"<li>{missing_tools} VMs ohne VMware Tools</li>")
                
                if 'snapshots' in report_data:
                    old_snapshots = sum(1 for snap in report_data['snapshots'] if snap.get('days_old', 0) > 30)
                    total_snap_size = sum(snap.get('size_gb', 0) for snap in report_data['snapshots'])
                    f.write(f"<li>{len(report_data['snapshots'])} aktive Snapshots, davon {old_snapshots} älter als 30 Tage</li>")
                    f.write(f"<li>Gesamtgröße aller Snapshots: {total_snap_size:.2f} GB</li>")
                
                if 'orphaned_vmdks' in report_data:
                    total_orphaned_size = sum(vmdk.get('size_gb', 0) for vmdk in report_data['orphaned_vmdks'])
                    f.write(f"<li>{len(report_data['orphaned_vmdks'])} verwaiste VMDK-Dateien mit insgesamt {total_orphaned_size:.2f} GB</li>")
                
                f.write("""
            </ul>
        </div>
    </div>
    
    <div class="footer mt-4">
        <div class="container">
            <div class="row">
                <div class="col-md-6">
                    <p>Erstellt mit VMware vSphere Reporter v29.0</p>
                </div>
                <div class="col-md-6 text-end">
                    <p>Copyright © 2025 Bechtle GmbH</p>
                </div>
            </div>
        </div>
    </div>
    
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
""")
            
            logger.info(f"HTML-Bericht erfolgreich generiert: {file_path}")
            return file_path
        
        except Exception as e:
            logger.error(f"Fehler beim Generieren des HTML-Berichts: {str(e)}")
            raise
    
    @handle_vsphere_errors
    def generate_pdf_report(self, filename, report_data):
        """
        Generiert einen PDF-Bericht.
        
        Args:
            filename (str): Name der Ausgabedatei
            report_data (dict): Daten für den Bericht
                
        Returns:
            str: Pfad zur generierten PDF-Datei
        """
        logger.info(f"Generiere PDF-Bericht: {filename}")
        
        try:
            # Verzeichnispfad erstellen
            file_path = os.path.join(self.output_dir, filename)
            
            # Dokument erstellen
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # Eigene Stile erstellen
            styles.add(ParagraphStyle(
                name='BechTitle',
                parent=styles['Heading1'],
                fontName='Helvetica-Bold',
                fontSize=18,
                textColor=colors.HexColor('#00355e'),
                spaceAfter=12
            ))
            
            styles.add(ParagraphStyle(
                name='BechHeading2',
                parent=styles['Heading2'],
                fontName='Helvetica-Bold',
                fontSize=14,
                textColor=colors.HexColor('#00355e'),
                spaceAfter=8
            ))
            
            styles.add(ParagraphStyle(
                name='BechNormal',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=10,
                spaceAfter=6
            ))
            
            styles.add(ParagraphStyle(
                name='BechInfo',
                parent=styles['Normal'],
                fontName='Helvetica-Oblique',
                fontSize=9,
                textColor=colors.gray
            ))
            
            # Story-Liste für alle Elemente
            story = []
            
            # Titelseite
            story.append(Paragraph(f"VMware vSphere Umgebungsbericht", styles['BechTitle']))
            story.append(Paragraph(f"Erstellt am: {self.timestamp}", styles['BechNormal']))
            story.append(Paragraph(f"vCenter Server: {self.vcenter_server}", styles['BechNormal']))
            story.append(Paragraph(f"Benutzer: {self.username}", styles['BechNormal']))
            if self.demo_mode:
                story.append(Paragraph("Demo-Modus aktiv", styles['BechInfo']))
            
            story.append(Spacer(1, 20))
            
            # VMware Tools Bericht
            if 'vmware_tools' in report_data and report_data['vmware_tools']:
                vmware_tools = report_data['vmware_tools']
                
                story.append(Paragraph("VMware Tools Status", styles['BechHeading2']))
                story.append(Paragraph(
                    "VMware Tools-Status für alle virtuellen Maschinen. "
                    "VMs mit veralteten Tools sind möglicherweise nicht kompatibel mit neuen Funktionen.",
                    styles['BechInfo']
                ))
                
                # Tabellendaten vorbereiten
                table_data = [['VM-Name', 'Status', 'Version', 'Laufzeitstatus', 'VM-Power']]
                
                # Sortieren nach Status (Problem first)
                sorted_vms = sorted(vmware_tools, key=lambda x: (
                    0 if x.get('tools_status') == 'NotInstalled' else
                    1 if x.get('tools_status') == 'UpdateNeeded' else
                    2 if x.get('tools_status') == 'NotRunning' else
                    3 if x.get('tools_status') == 'Unmanaged' else 4
                ))
                
                for vm in sorted_vms:
                    status_text = ""
                    if vm.get('tools_status') == 'Current':
                        status_text = "Aktuell"
                    elif vm.get('tools_status') == 'UpdateNeeded':
                        status_text = "Update verfügbar"
                    elif vm.get('tools_status') == 'NotInstalled':
                        status_text = "Nicht installiert"
                    elif vm.get('tools_status') == 'NotRunning':
                        status_text = "Nicht ausgeführt"
                    else:
                        status_text = "Unbekannt"
                    
                    power = 'Eingeschaltet' if vm.get('power_state') == 'poweredOn' else 'Ausgeschaltet'
                    
                    table_data.append([
                        vm.get('vm_name', 'Unbekannt'),
                        status_text,
                        vm.get('tools_version', 'Unbekannt'),
                        vm.get('tools_running_status', 'Unbekannt'),
                        power
                    ])
                
                # Tabelle erstellen
                table = Table(table_data, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f3f3')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#00355e')),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('TOPPADDING', (0, 1), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ]))
                
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Snapshots Bericht
            if 'snapshots' in report_data and report_data['snapshots']:
                snapshots = report_data['snapshots']
                
                story.append(Paragraph("VM-Snapshots", styles['BechHeading2']))
                story.append(Paragraph(
                    "Liste aller aktiven Snapshots in der Umgebung. "
                    "Ältere Snapshots könnten zu Leistungsproblemen führen oder Speicherplatz vergeuden.",
                    styles['BechInfo']
                ))
                
                # Tabellendaten vorbereiten
                table_data = [['VM-Name', 'Snapshot-Name', 'Erstellt am', 'Alter (Tage)', 'Größe']]
                
                # Sortieren nach Alter (älteste zuerst)
                sorted_snapshots = sorted(snapshots, key=lambda x: x.get('days_old', 0), reverse=True)
                
                for snapshot in sorted_snapshots:
                    table_data.append([
                        snapshot.get('vm_name', 'Unbekannt'),
                        snapshot.get('name', 'Unbekannt'),
                        snapshot.get('created_date', 'Unbekannt'),
                        str(snapshot.get('days_old', 'Unbekannt')),
                        f"{snapshot.get('size_gb', 'Unbekannt')} GB"
                    ])
                
                # Tabelle erstellen
                table = Table(table_data, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f3f3')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#00355e')),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('TOPPADDING', (0, 1), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ]))
                
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Verwaiste VMDKs Bericht
            if 'orphaned_vmdks' in report_data and report_data['orphaned_vmdks']:
                orphaned_vmdks = report_data['orphaned_vmdks']
                
                story.append(Paragraph("Verwaiste VMDK-Dateien", styles['BechHeading2']))
                story.append(Paragraph(
                    "Liste aller verwaisten VMDK-Dateien in der Umgebung. "
                    "Diese belegen Speicherplatz, werden aber von keiner VM verwendet.",
                    styles['BechInfo']
                ))
                
                # Tabellendaten vorbereiten
                table_data = [['Datastore', 'Pfad', 'Größe (GB)', 'Alter (Tage)', 'Empfohlene Aktion']]
                
                # Sortieren nach Größe (größte zuerst)
                sorted_vmdks = sorted(orphaned_vmdks, key=lambda x: x.get('size_gb', 0), reverse=True)
                
                for vmdk in sorted_vmdks:
                    table_data.append([
                        vmdk.get('datastore', 'Unbekannt'),
                        vmdk.get('path', 'Unbekannt'),
                        str(vmdk.get('size_gb', 'Unbekannt')),
                        str(vmdk.get('days_orphaned', 'Unbekannt')),
                        vmdk.get('recovery_action', 'Unbekannt')
                    ])
                
                # Tabelle erstellen
                table = Table(table_data, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f3f3')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#00355e')),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('TOPPADDING', (0, 1), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ]))
                
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Zusammenfassung
            story.append(Paragraph("Zusammenfassung", styles['BechHeading2']))
            
            summary_text = []
            
            if 'vmware_tools' in report_data:
                outdated_tools = sum(1 for vm in report_data['vmware_tools'] if vm.get('tools_status') == 'UpdateNeeded')
                missing_tools = sum(1 for vm in report_data['vmware_tools'] if vm.get('tools_status') == 'NotInstalled')
                summary_text.append(f"{outdated_tools} VMs mit veralteten VMware Tools")
                summary_text.append(f"{missing_tools} VMs ohne VMware Tools")
            
            if 'snapshots' in report_data:
                old_snapshots = sum(1 for snap in report_data['snapshots'] if snap.get('days_old', 0) > 30)
                total_snap_size = sum(snap.get('size_gb', 0) for snap in report_data['snapshots'])
                summary_text.append(f"{len(report_data['snapshots'])} aktive Snapshots, davon {old_snapshots} älter als 30 Tage")
                summary_text.append(f"Gesamtgröße aller Snapshots: {total_snap_size:.2f} GB")
            
            if 'orphaned_vmdks' in report_data:
                total_orphaned_size = sum(vmdk.get('size_gb', 0) for vmdk in report_data['orphaned_vmdks'])
                summary_text.append(f"{len(report_data['orphaned_vmdks'])} verwaiste VMDK-Dateien mit insgesamt {total_orphaned_size:.2f} GB")
            
            for line in summary_text:
                story.append(Paragraph(f"• {line}", styles['BechNormal']))
            
            # Footer
            story.append(Spacer(1, 40))
            story.append(Paragraph(f"Erstellt mit VMware vSphere Reporter v29.0", styles['BechInfo']))
            story.append(Paragraph(f"Copyright © 2025 Bechtle GmbH", styles['BechInfo']))
            
            # PDF generieren
            doc.build(story)
            
            logger.info(f"PDF-Bericht erfolgreich generiert: {file_path}")
            return file_path
        
        except Exception as e:
            logger.error(f"Fehler beim Generieren des PDF-Berichts: {str(e)}")
            raise
    
    @handle_vsphere_errors
    def generate_docx_report(self, filename, report_data):
        """
        Generiert einen DOCX-Bericht.
        
        Args:
            filename (str): Name der Ausgabedatei
            report_data (dict): Daten für den Bericht
                
        Returns:
            str: Pfad zur generierten DOCX-Datei
        """
        logger.info(f"Generiere DOCX-Bericht: {filename}")
        
        try:
            # DOCX-Modul importieren
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL
            
            # Verzeichnispfad erstellen
            file_path = os.path.join(self.output_dir, filename)
            
            # Dokument erstellen
            doc = Document()
            
            # Ränder festlegen
            for section in doc.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
            
            # Titel und Kopfinformationen
            title = doc.add_heading('VMware vSphere Umgebungsbericht', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f"Erstellt am: {self.timestamp}")
            doc.add_paragraph(f"vCenter Server: {self.vcenter_server}")
            doc.add_paragraph(f"Benutzer: {self.username}")
            
            if self.demo_mode:
                p = doc.add_paragraph()
                run = p.add_run("Demo-Modus aktiv")
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()
            
            # VMware Tools Bericht
            if 'vmware_tools' in report_data and report_data['vmware_tools']:
                vmware_tools = report_data['vmware_tools']
                
                doc.add_heading('VMware Tools Status', 1)
                
                p = doc.add_paragraph("VMware Tools-Status für alle virtuellen Maschinen. ")
                p.add_run("VMs mit veralteten Tools sind möglicherweise nicht kompatibel mit neuen Funktionen.")
                
                # Tabelle erstellen
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Spaltenüberschriften
                header_cells = table.rows[0].cells
                header_cells[0].text = 'VM-Name'
                header_cells[1].text = 'Status'
                header_cells[2].text = 'Version'
                header_cells[3].text = 'Laufzeitstatus'
                header_cells[4].text = 'VM-Power'
                
                # Überschriften formatieren
                for cell in header_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                
                # Sortieren nach Status (Problem first)
                sorted_vms = sorted(vmware_tools, key=lambda x: (
                    0 if x.get('tools_status') == 'NotInstalled' else
                    1 if x.get('tools_status') == 'UpdateNeeded' else
                    2 if x.get('tools_status') == 'NotRunning' else
                    3 if x.get('tools_status') == 'Unmanaged' else 4
                ))
                
                # Zeilen hinzufügen
                for vm in sorted_vms:
                    status_text = ""
                    if vm.get('tools_status') == 'Current':
                        status_text = "Aktuell"
                    elif vm.get('tools_status') == 'UpdateNeeded':
                        status_text = "Update verfügbar"
                    elif vm.get('tools_status') == 'NotInstalled':
                        status_text = "Nicht installiert"
                    elif vm.get('tools_status') == 'NotRunning':
                        status_text = "Nicht ausgeführt"
                    else:
                        status_text = "Unbekannt"
                    
                    power = 'Eingeschaltet' if vm.get('power_state') == 'poweredOn' else 'Ausgeschaltet'
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = vm.get('vm_name', 'Unbekannt')
                    row_cells[1].text = status_text
                    row_cells[2].text = vm.get('tools_version', 'Unbekannt')
                    row_cells[3].text = vm.get('tools_running_status', 'Unbekannt')
                    row_cells[4].text = power
                    
                    # Zellen vertikal zentrieren
                    for cell in row_cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                doc.add_paragraph()
            
            # Snapshots Bericht
            if 'snapshots' in report_data and report_data['snapshots']:
                snapshots = report_data['snapshots']
                
                doc.add_heading('VM-Snapshots', 1)
                
                p = doc.add_paragraph("Liste aller aktiven Snapshots in der Umgebung. ")
                p.add_run("Ältere Snapshots könnten zu Leistungsproblemen führen oder Speicherplatz vergeuden.")
                
                # Tabelle erstellen
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Spaltenüberschriften
                header_cells = table.rows[0].cells
                header_cells[0].text = 'VM-Name'
                header_cells[1].text = 'Snapshot-Name'
                header_cells[2].text = 'Erstellt am'
                header_cells[3].text = 'Alter (Tage)'
                header_cells[4].text = 'Größe'
                
                # Überschriften formatieren
                for cell in header_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                
                # Sortieren nach Alter (älteste zuerst)
                sorted_snapshots = sorted(snapshots, key=lambda x: x.get('days_old', 0), reverse=True)
                
                # Zeilen hinzufügen
                for snapshot in sorted_snapshots:
                    row_cells = table.add_row().cells
                    row_cells[0].text = snapshot.get('vm_name', 'Unbekannt')
                    row_cells[1].text = snapshot.get('name', 'Unbekannt')
                    row_cells[2].text = snapshot.get('created_date', 'Unbekannt')
                    row_cells[3].text = str(snapshot.get('days_old', 'Unbekannt'))
                    row_cells[4].text = f"{snapshot.get('size_gb', 'Unbekannt')} GB"
                    
                    # Zellen vertikal zentrieren
                    for cell in row_cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                doc.add_paragraph()
            
            # Verwaiste VMDKs Bericht
            if 'orphaned_vmdks' in report_data and report_data['orphaned_vmdks']:
                orphaned_vmdks = report_data['orphaned_vmdks']
                
                doc.add_heading('Verwaiste VMDK-Dateien', 1)
                
                p = doc.add_paragraph("Liste aller verwaisten VMDK-Dateien in der Umgebung. ")
                p.add_run("Diese belegen Speicherplatz, werden aber von keiner VM verwendet.")
                
                # Tabelle erstellen
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Spaltenüberschriften
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Datastore'
                header_cells[1].text = 'Pfad'
                header_cells[2].text = 'Größe (GB)'
                header_cells[3].text = 'Alter (Tage)'
                header_cells[4].text = 'Empfohlene Aktion'
                
                # Überschriften formatieren
                for cell in header_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                
                # Sortieren nach Größe (größte zuerst)
                sorted_vmdks = sorted(orphaned_vmdks, key=lambda x: x.get('size_gb', 0), reverse=True)
                
                # Zeilen hinzufügen
                for vmdk in sorted_vmdks:
                    row_cells = table.add_row().cells
                    row_cells[0].text = vmdk.get('datastore', 'Unbekannt')
                    row_cells[1].text = vmdk.get('path', 'Unbekannt')
                    row_cells[2].text = str(vmdk.get('size_gb', 'Unbekannt'))
                    row_cells[3].text = str(vmdk.get('days_orphaned', 'Unbekannt'))
                    row_cells[4].text = vmdk.get('recovery_action', 'Unbekannt')
                    
                    # Zellen vertikal zentrieren
                    for cell in row_cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                doc.add_paragraph()
            
            # Zusammenfassung
            doc.add_heading('Zusammenfassung', 1)
            
            # Aufzählungsliste erstellen
            if 'vmware_tools' in report_data:
                outdated_tools = sum(1 for vm in report_data['vmware_tools'] if vm.get('tools_status') == 'UpdateNeeded')
                missing_tools = sum(1 for vm in report_data['vmware_tools'] if vm.get('tools_status') == 'NotInstalled')
                p = doc.add_paragraph(f"{outdated_tools} VMs mit veralteten VMware Tools", style='List Bullet')
                p = doc.add_paragraph(f"{missing_tools} VMs ohne VMware Tools", style='List Bullet')
            
            if 'snapshots' in report_data:
                old_snapshots = sum(1 for snap in report_data['snapshots'] if snap.get('days_old', 0) > 30)
                total_snap_size = sum(snap.get('size_gb', 0) for snap in report_data['snapshots'])
                p = doc.add_paragraph(f"{len(report_data['snapshots'])} aktive Snapshots, davon {old_snapshots} älter als 30 Tage", style='List Bullet')
                p = doc.add_paragraph(f"Gesamtgröße aller Snapshots: {total_snap_size:.2f} GB", style='List Bullet')
            
            if 'orphaned_vmdks' in report_data:
                total_orphaned_size = sum(vmdk.get('size_gb', 0) for vmdk in report_data['orphaned_vmdks'])
                p = doc.add_paragraph(f"{len(report_data['orphaned_vmdks'])} verwaiste VMDK-Dateien mit insgesamt {total_orphaned_size:.2f} GB", style='List Bullet')
            
            # Footer
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Erstellt mit VMware vSphere Reporter v29.0 | Copyright © 2025 Bechtle GmbH"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Dokument speichern
            doc.save(file_path)
            
            logger.info(f"DOCX-Bericht erfolgreich generiert: {file_path}")
            return file_path
        
        except Exception as e:
            logger.error(f"Fehler beim Generieren des DOCX-Berichts: {str(e)}")
            raise
        
        # Vollständigen Pfad zur Ausgabedatei erstellen
        output_path = os.path.join(self.output_dir, filename)
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # Kopfbereich des HTML-Dokuments
                f.write(f"""<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>VMware vSphere Report - {self.datestamp}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            margin: 0;
            padding: 0;
            color: #333;
            background-color: #f8f9fa;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
        }}
        .header {{
            background-color: {BECHTLE_DARK_BLUE};
            color: white;
            padding: 20px;
            margin-bottom: 30px;
            border-radius: 5px;
        }}
        .header h1 {{
            margin: 0;
            font-size: 24px;
        }}
        .header p {{
            margin: 5px 0 0 0;
            font-size: 14px;
            opacity: 0.8;
        }}
        .section {{
            background-color: white;
            border-radius: 5px;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            margin-bottom: 30px;
            overflow: hidden;
        }}
        .section-header {{
            background-color: {BECHTLE_DARK_BLUE};
            color: white;
            padding: 15px 20px;
        }}
        .section-header h2 {{
            margin: 0;
            font-size: 18px;
        }}
        .section-content {{
            padding: 20px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
        }}
        th {{
            background-color: {BECHTLE_DARK_BLUE};
            color: white;
            padding: 10px;
            text-align: left;
        }}
        td {{
            padding: 10px;
            border-bottom: 1px solid #eee;
        }}
        tr:nth-child(even) {{
            background-color: {BECHTLE_LIGHT_GRAY};
        }}
        .footer {{
            text-align: center;
            margin-top: 30px;
            color: #666;
            font-size: 12px;
        }}
        .status-ok {{
            color: {BECHTLE_GREEN};
            font-weight: bold;
        }}
        .status-warning {{
            color: {BECHTLE_ORANGE};
            font-weight: bold;
        }}
        .status-error {{
            color: #dc3545;
            font-weight: bold;
        }}
        .age-recent {{
            color: {BECHTLE_GREEN};
        }}
        .age-warning {{
            color: {BECHTLE_ORANGE};
        }}
        .age-danger {{
            color: #dc3545;
        }}
        .demo-badge {{
            display: inline-block;
            background-color: {BECHTLE_ORANGE};
            color: white;
            padding: 3px 8px;
            border-radius: 3px;
            font-size: 12px;
            margin-left: 10px;
            vertical-align: middle;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>VMware vSphere Infrastruktur-Bericht {' <span class="demo-badge">Demo-Modus</span>' if self.demo_mode else ''}</h1>
            <p>Erstellt am {self.timestamp} für {self.vcenter_server} von {self.username}</p>
        </div>
""")
                
                # VMware Tools-Abschnitt
                if report_data.get('vmware_tools'):
                    vmware_tools_data = report_data['vmware_tools']
                    f.write("""
        <div class="section">
            <div class="section-header">
                <h2>VMware Tools Status</h2>
            </div>
            <div class="section-content">
                <table>
                    <thead>
                        <tr>
                            <th>VM-Name</th>
                            <th>ESXi-Host</th>
                            <th>Tools-Version</th>
                            <th>Status</th>
                            <th>Betriebssystem</th>
                            <th>Letzter Boot</th>
                        </tr>
                    </thead>
                    <tbody>
""")
                    
                    for vm in vmware_tools_data:
                        # Status-Klasse bestimmen
                        status_class = "status-ok"
                        status_text = "Aktuell"
                        
                        if vm['tools_status'] == 'NotInstalled':
                            status_class = "status-error"
                            status_text = "Nicht installiert"
                        elif vm['tools_status'] == 'UpdateNeeded':
                            status_class = "status-warning"
                            status_text = "Update erforderlich"
                        elif vm['tools_status'] == 'NotRunning':
                            status_class = "status-warning"
                            status_text = "Nicht ausgeführt"
                        elif vm['tools_status'] == 'Unmanaged':
                            status_class = ""
                            status_text = "Nicht verwaltet"
                        
                        f.write(f"""
                        <tr>
                            <td>{vm['vm_name']}</td>
                            <td>{vm['esxi_host']}</td>
                            <td>{vm['tools_version']}</td>
                            <td class="{status_class}">{status_text}</td>
                            <td>{vm['os']}</td>
                            <td>{vm['last_boot']}</td>
                        </tr>
""")
                    
                    f.write("""
                    </tbody>
                </table>
            </div>
        </div>
""")
                
                # Snapshots-Abschnitt
                if report_data.get('snapshots'):
                    snapshots_data = report_data['snapshots']
                    f.write("""
        <div class="section">
            <div class="section-header">
                <h2>VM-Snapshots</h2>
            </div>
            <div class="section-content">
                <table>
                    <thead>
                        <tr>
                            <th>VM-Name</th>
                            <th>Snapshot-Name</th>
                            <th>Beschreibung</th>
                            <th>Erstellt am</th>
                            <th>Alter (Tage)</th>
                            <th>Größe (GB)</th>
                            <th>ESXi-Host</th>
                        </tr>
                    </thead>
                    <tbody>
""")
                    
                    for snapshot in snapshots_data:
                        # Alterskategorie bestimmen
                        age_class = snapshot.get('age_category', '')
                        if age_class == 'recent':
                            age_class = "age-recent"
                        elif age_class == 'warning':
                            age_class = "age-warning"
                        elif age_class == 'danger':
                            age_class = "age-danger"
                        
                        f.write(f"""
                        <tr>
                            <td>{snapshot['vm_name']}</td>
                            <td>{snapshot['snapshot_name']}</td>
                            <td>{snapshot['description']}</td>
                            <td>{snapshot['date_created']}</td>
                            <td class="{age_class}">{snapshot['days_old']}</td>
                            <td>{snapshot['size_gb']}</td>
                            <td>{snapshot['esxi_host']}</td>
                        </tr>
""")
                    
                    f.write("""
                    </tbody>
                </table>
            </div>
        </div>
""")
                
                # Verwaiste VMDKs-Abschnitt
                if report_data.get('orphaned_vmdks'):
                    orphaned_vmdks_data = report_data['orphaned_vmdks']
                    f.write("""
        <div class="section">
            <div class="section-header">
                <h2>Verwaiste VMDK-Dateien</h2>
            </div>
            <div class="section-content">
                <table>
                    <thead>
                        <tr>
                            <th>Datastore</th>
                            <th>Pfad</th>
                            <th>Größe (GB)</th>
                            <th>Letzte Änderung</th>
                            <th>Tage verwaist</th>
                            <th>Thin Provisioned</th>
                            <th>Vermutliche VM</th>
                            <th>Empfohlene Aktion</th>
                        </tr>
                    </thead>
                    <tbody>
""")
                    
                    for vmdk in orphaned_vmdks_data:
                        # Altersklasse bestimmen
                        age_class = ""
                        if vmdk['days_orphaned'] > 90:
                            age_class = "age-danger"
                        elif vmdk['days_orphaned'] > 30:
                            age_class = "age-warning"
                        
                        thin_provisioned = "Ja" if vmdk.get('thin_provisioned', False) else "Nein"
                        
                        f.write(f"""
                        <tr>
                            <td>{vmdk['datastore']}</td>
                            <td>{vmdk['path']}</td>
                            <td>{vmdk['size_gb']}</td>
                            <td>{vmdk['last_modified']}</td>
                            <td class="{age_class}">{vmdk['days_orphaned']}</td>
                            <td>{thin_provisioned}</td>
                            <td>{vmdk['probable_vm']}</td>
                            <td>{vmdk['recovery_action']}</td>
                        </tr>
""")
                    
                    f.write("""
                    </tbody>
                </table>
            </div>
        </div>
""")
                
                # Fußzeile und Abschluss des HTML-Dokuments
                f.write(f"""
        <div class="footer">
            <p>VMware vSphere Reporter v29.0 | Copyright &copy; 2025 Bechtle GmbH | Alle Rechte vorbehalten.</p>
            <p>VMware ist eine eingetragene Marke der VMware, Inc.</p>
        </div>
    </div>
</body>
</html>
""")
            
            logger.info(f"HTML-Bericht erfolgreich generiert: {output_path}")
            return output_path
            
        except Exception as e:
            error_message = f"Fehler bei der Generierung des HTML-Berichts: {str(e)}"
            logger.error(error_message)
            raise VSphereReportGenerationError(error_message)
    
    @handle_vsphere_errors
    def generate_pdf_report(self, filename, report_data):
        """
        Generiert einen PDF-Bericht.
        
        Args:
            filename (str): Name der Ausgabedatei
            report_data (dict): Daten für den Bericht
                
        Returns:
            str: Pfad zur generierten PDF-Datei
        """
        logger.info(f"Generiere PDF-Bericht: {filename}")
        
        # Vollständigen Pfad zur Ausgabedatei erstellen
        output_path = os.path.join(self.output_dir, filename)
        
        try:
            # Stelle sicher, dass die ReportLab-Bibliothek verfügbar ist
            if 'reportlab' not in sys.modules:
                raise ImportError("ReportLab-Bibliothek ist nicht verfügbar.")
            
            # Erstelle PDF-Dokument
            doc = SimpleDocTemplate(output_path, pagesize=A4, 
                                 rightMargin=72, leftMargin=72,
                                 topMargin=72, bottomMargin=72)
            
            # Definiere Stile
            styles = getSampleStyleSheet()
            title_style = styles['Heading1']
            title_style.alignment = TA_CENTER
            title_style.textColor = colors.HexColor(BECHTLE_DARK_BLUE)
            
            section_style = styles['Heading2']
            section_style.textColor = colors.HexColor(BECHTLE_DARK_BLUE)
            
            normal_style = styles['Normal']
            
            # Liste für PDF-Elemente
            elements = []
            
            # Titel
            title_text = "VMware vSphere Infrastruktur-Bericht"
            if self.demo_mode:
                title_text += " (Demo-Modus)"
            elements.append(Paragraph(title_text, title_style))
            elements.append(Spacer(1, 12))
            
            # Unterüberschrift
            subtitle_text = f"Erstellt am {self.timestamp} für {self.vcenter_server} von {self.username}"
            elements.append(Paragraph(subtitle_text, normal_style))
            elements.append(Spacer(1, 20))
            
            # VMware Tools-Abschnitt
            if report_data.get('vmware_tools'):
                vmware_tools_data = report_data['vmware_tools']
                elements.append(Paragraph("VMware Tools Status", section_style))
                elements.append(Spacer(1, 12))
                
                # Tabellendaten vorbereiten
                table_data = [["VM-Name", "ESXi-Host", "Tools-Version", "Status", "Betriebssystem", "Letzter Boot"]]
                
                for vm in vmware_tools_data:
                    status_text = "Aktuell"
                    if vm['tools_status'] == 'NotInstalled':
                        status_text = "Nicht installiert"
                    elif vm['tools_status'] == 'UpdateNeeded':
                        status_text = "Update erforderlich"
                    elif vm['tools_status'] == 'NotRunning':
                        status_text = "Nicht ausgeführt"
                    elif vm['tools_status'] == 'Unmanaged':
                        status_text = "Nicht verwaltet"
                    
                    table_data.append([
                        vm['vm_name'],
                        vm['esxi_host'],
                        vm['tools_version'],
                        status_text,
                        vm['os'],
                        vm['last_boot']
                    ])
                
                # Tabelle erstellen
                table = Table(table_data, repeatRows=1)
                
                # Tabellenstil
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BECHTLE_DARK_BLUE)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey)
                ])
                
                # Abwechselnde Zeilenfarben
                for i in range(1, len(table_data), 2):
                    table_style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor(BECHTLE_LIGHT_GRAY))
                
                table.setStyle(table_style)
                elements.append(table)
                elements.append(Spacer(1, 20))
            
            # Snapshots-Abschnitt
            if report_data.get('snapshots'):
                snapshots_data = report_data['snapshots']
                elements.append(Paragraph("VM-Snapshots", section_style))
                elements.append(Spacer(1, 12))
                
                # Tabellendaten vorbereiten
                table_data = [["VM-Name", "Snapshot-Name", "Erstellt am", "Alter (Tage)", "Größe (GB)", "ESXi-Host"]]
                
                for snapshot in snapshots_data:
                    table_data.append([
                        snapshot['vm_name'],
                        snapshot['snapshot_name'],
                        snapshot['date_created'],
                        snapshot['days_old'],
                        snapshot['size_gb'],
                        snapshot['esxi_host']
                    ])
                
                # Tabelle erstellen
                table = Table(table_data, repeatRows=1)
                
                # Tabellenstil
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BECHTLE_DARK_BLUE)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey)
                ])
                
                # Abwechselnde Zeilenfarben
                for i in range(1, len(table_data), 2):
                    table_style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor(BECHTLE_LIGHT_GRAY))
                
                table.setStyle(table_style)
                elements.append(table)
                elements.append(Spacer(1, 20))
            
            # Verwaiste VMDKs-Abschnitt
            if report_data.get('orphaned_vmdks'):
                orphaned_vmdks_data = report_data['orphaned_vmdks']
                elements.append(Paragraph("Verwaiste VMDK-Dateien", section_style))
                elements.append(Spacer(1, 12))
                
                # Tabellendaten vorbereiten
                table_data = [["Datastore", "Pfad", "Größe (GB)", "Letzte Änderung", "Tage verwaist", "Empfohlene Aktion"]]
                
                for vmdk in orphaned_vmdks_data:
                    thin_provisioned = "Ja" if vmdk.get('thin_provisioned', False) else "Nein"
                    
                    table_data.append([
                        vmdk['datastore'],
                        vmdk['path'],
                        vmdk['size_gb'],
                        vmdk['last_modified'],
                        vmdk['days_orphaned'],
                        vmdk['recovery_action']
                    ])
                
                # Tabelle erstellen
                table = Table(table_data, repeatRows=1)
                
                # Tabellenstil
                table_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BECHTLE_DARK_BLUE)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey)
                ])
                
                # Abwechselnde Zeilenfarben
                for i in range(1, len(table_data), 2):
                    table_style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor(BECHTLE_LIGHT_GRAY))
                
                table.setStyle(table_style)
                elements.append(table)
                elements.append(Spacer(1, 20))
            
            # Fußzeile
            footer_text = "VMware vSphere Reporter v29.0 | Copyright © 2025 Bechtle GmbH | Alle Rechte vorbehalten.\n"
            footer_text += "VMware ist eine eingetragene Marke der VMware, Inc."
            
            footer_style = ParagraphStyle(
                name='FooterStyle',
                parent=styles['Normal'],
                fontSize=8,
                textColor=colors.gray,
                alignment=TA_CENTER
            )
            
            elements.append(Paragraph(footer_text, footer_style))
            
            # PDF erstellen
            doc.build(elements)
            
            logger.info(f"PDF-Bericht erfolgreich generiert: {output_path}")
            return output_path
            
        except ImportError as e:
            error_message = f"ReportLab-Bibliothek fehlt für PDF-Erstellung: {str(e)}"
            logger.error(error_message)
            raise VSphereReportGenerationError(error_message)
        except Exception as e:
            error_message = f"Fehler bei der Generierung des PDF-Berichts: {str(e)}"
            logger.error(error_message)
            raise VSphereReportGenerationError(error_message)
    
    @handle_vsphere_errors
    def generate_docx_report(self, filename, report_data):
        """
        Generiert einen DOCX-Bericht.
        
        Args:
            filename (str): Name der Ausgabedatei
            report_data (dict): Daten für den Bericht
                
        Returns:
            str: Pfad zur generierten DOCX-Datei
        """
        logger.info(f"Generiere DOCX-Bericht: {filename}")
        
        # Vollständigen Pfad zur Ausgabedatei erstellen
        output_path = os.path.join(self.output_dir, filename)
        
        try:
            # Stelle sicher, dass die python-docx-Bibliothek verfügbar ist
            if 'docx' not in sys.modules:
                raise ImportError("python-docx-Bibliothek ist nicht verfügbar.")
            
            # Erstelle Word-Dokument
            doc = docx.Document()
            
            # Seitenränder anpassen
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Titel
            title = doc.add_heading("VMware vSphere Infrastruktur-Bericht", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if self.demo_mode:
                title.add_run(" (Demo-Modus)")
            
            # Unterüberschrift
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.add_run(f"Erstellt am {self.timestamp} für {self.vcenter_server} von {self.username}")
            
            # Trenner
            doc.add_paragraph()
            
            # VMware Tools-Abschnitt
            if report_data.get('vmware_tools'):
                vmware_tools_data = report_data['vmware_tools']
                doc.add_heading("VMware Tools Status", level=2)
                
                # Tabelle erstellen
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                table.autofit = False
                
                # Überschriften
                header_cells = table.rows[0].cells
                header_cells[0].text = "VM-Name"
                header_cells[1].text = "ESXi-Host"
                header_cells[2].text = "Tools-Version"
                header_cells[3].text = "Status"
                header_cells[4].text = "Betriebssystem"
                header_cells[5].text = "Letzter Boot"
                
                # Formatiere Überschriften fett auf blauem Hintergrund
                for cell in header_cells:
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    shade_table_cell(cell, BECHTLE_DARK_BLUE)
                
                # Daten einfügen
                for vm in vmware_tools_data:
                    # Status-Text und Farbe bestimmen
                    status_text = "Aktuell"
                    status_color = BECHTLE_GREEN
                    
                    if vm['tools_status'] == 'NotInstalled':
                        status_text = "Nicht installiert"
                        status_color = "DC3545"  # Rot
                    elif vm['tools_status'] == 'UpdateNeeded':
                        status_text = "Update erforderlich"
                        status_color = BECHTLE_ORANGE
                    elif vm['tools_status'] == 'NotRunning':
                        status_text = "Nicht ausgeführt"
                        status_color = BECHTLE_ORANGE
                    elif vm['tools_status'] == 'Unmanaged':
                        status_text = "Nicht verwaltet"
                        status_color = BECHTLE_DARK_GRAY
                    
                    # Neue Zeile hinzufügen
                    row_cells = table.add_row().cells
                    row_cells[0].text = vm['vm_name']
                    row_cells[1].text = vm['esxi_host']
                    row_cells[2].text = vm['tools_version']
                    row_cells[3].text = status_text
                    
                    # Status-Farbe anwenden
                    status_paragraph = row_cells[3].paragraphs[0]
                    status_run = status_paragraph.runs[0]
                    status_run.font.color.rgb = RGBColor.from_string(status_color)
                    status_run.font.bold = True
                    
                    row_cells[4].text = vm['os']
                    row_cells[5].text = vm['last_boot']
                
                doc.add_paragraph()
            
            # Snapshots-Abschnitt
            if report_data.get('snapshots'):
                snapshots_data = report_data['snapshots']
                doc.add_heading("VM-Snapshots", level=2)
                
                # Tabelle erstellen
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'
                table.autofit = False
                
                # Überschriften
                header_cells = table.rows[0].cells
                header_cells[0].text = "VM-Name"
                header_cells[1].text = "Snapshot-Name"
                header_cells[2].text = "Beschreibung"
                header_cells[3].text = "Erstellt am"
                header_cells[4].text = "Alter (Tage)"
                header_cells[5].text = "Größe (GB)"
                header_cells[6].text = "ESXi-Host"
                
                # Formatiere Überschriften fett auf blauem Hintergrund
                for cell in header_cells:
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    shade_table_cell(cell, BECHTLE_DARK_BLUE)
                
                # Daten einfügen
                for snapshot in snapshots_data:
                    # Altersfarbe bestimmen
                    age_color = BECHTLE_GREEN
                    if snapshot.get('age_category') == 'warning':
                        age_color = BECHTLE_ORANGE
                    elif snapshot.get('age_category') == 'danger':
                        age_color = "DC3545"  # Rot
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = snapshot['vm_name']
                    row_cells[1].text = snapshot['snapshot_name']
                    row_cells[2].text = snapshot['description']
                    row_cells[3].text = snapshot['date_created']
                    row_cells[4].text = str(snapshot['days_old'])
                    
                    # Altersfarbe anwenden
                    age_paragraph = row_cells[4].paragraphs[0]
                    age_run = age_paragraph.runs[0]
                    age_run.font.color.rgb = RGBColor.from_string(age_color)
                    if snapshot.get('age_category') != 'recent':
                        age_run.font.bold = True
                    
                    row_cells[5].text = str(snapshot['size_gb'])
                    row_cells[6].text = snapshot['esxi_host']
                
                doc.add_paragraph()
            
            # Verwaiste VMDKs-Abschnitt
            if report_data.get('orphaned_vmdks'):
                orphaned_vmdks_data = report_data['orphaned_vmdks']
                doc.add_heading("Verwaiste VMDK-Dateien", level=2)
                
                # Tabelle erstellen
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                table.autofit = False
                
                # Überschriften
                header_cells = table.rows[0].cells
                header_cells[0].text = "Datastore"
                header_cells[1].text = "Pfad"
                header_cells[2].text = "Größe (GB)"
                header_cells[3].text = "Letzte Änderung"
                header_cells[4].text = "Tage verwaist"
                header_cells[5].text = "Thin Provisioned"
                header_cells[6].text = "Vermutliche VM"
                header_cells[7].text = "Empfohlene Aktion"
                
                # Formatiere Überschriften fett auf blauem Hintergrund
                for cell in header_cells:
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    shade_table_cell(cell, BECHTLE_DARK_BLUE)
                
                # Daten einfügen
                for vmdk in orphaned_vmdks_data:
                    # Altersfarbe bestimmen
                    age_color = "000000"  # Schwarz
                    if vmdk['days_orphaned'] > 90:
                        age_color = "DC3545"  # Rot
                    elif vmdk['days_orphaned'] > 30:
                        age_color = BECHTLE_ORANGE
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = vmdk['datastore']
                    row_cells[1].text = vmdk['path']
                    row_cells[2].text = str(vmdk['size_gb'])
                    row_cells[3].text = vmdk['last_modified']
                    row_cells[4].text = str(vmdk['days_orphaned'])
                    
                    # Altersfarbe anwenden
                    age_paragraph = row_cells[4].paragraphs[0]
                    age_run = age_paragraph.runs[0]
                    age_run.font.color.rgb = RGBColor.from_string(age_color)
                    if vmdk['days_orphaned'] > 30:
                        age_run.font.bold = True
                    
                    row_cells[5].text = "Ja" if vmdk.get('thin_provisioned', False) else "Nein"
                    row_cells[6].text = vmdk['probable_vm']
                    row_cells[7].text = vmdk['recovery_action']
                
                doc.add_paragraph()
            
            # Fußzeile
            footer = doc.sections[0].footer
            footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_text = "VMware vSphere Reporter v29.0 | Copyright © 2025 Bechtle GmbH | Alle Rechte vorbehalten. "
            footer_text += "VMware ist eine eingetragene Marke der VMware, Inc."
            footer_paragraph.text = footer_text
            
            # Word-Dokument speichern
            doc.save(output_path)
            
            logger.info(f"DOCX-Bericht erfolgreich generiert: {output_path}")
            return output_path
            
        except ImportError as e:
            error_message = f"python-docx-Bibliothek fehlt für DOCX-Erstellung: {str(e)}"
            logger.error(error_message)
            raise VSphereReportGenerationError(error_message)
        except Exception as e:
            error_message = f"Fehler bei der Generierung des DOCX-Berichts: {str(e)}"
            logger.error(error_message)
            raise VSphereReportGenerationError(error_message)

# Hilfsfunktion für Word-Tabellenzellen-Hintergrund
def shade_table_cell(cell, color_hex):
    """
    Setzt den Hintergrund einer Tabellenzelle in Word.
    
    Args:
        cell: Die Zelle
        color_hex: Die Farbe im Hex-Format (#RRGGBB)
    """
    color_hex = color_hex.lstrip('#')
    if len(color_hex) != 6:
        return
    
    try:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        
        shading = cell._element.get_or_add_tcPr()
        shading.append(docx.oxml.parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="{color_hex}"/>'))
    except:
        # Fallback if there's an error
        pass