#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
DOCX report exporter
"""

import os
import logging
import datetime
import humanize
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

class DOCXExporter:
    """Exporter for DOCX reports"""
    
    def __init__(self, data, timestamp):
        """
        Initialize the DOCX exporter
        
        Args:
            data (dict): Dictionary containing collected vSphere data
            timestamp (datetime): Report generation timestamp
        """
        self.data = data
        self.timestamp = timestamp
        self.document = Document()
        
    def export(self, output_path):
        """
        Export data to DOCX file
        
        Args:
            output_path (str): Path to save the DOCX file
            
        Returns:
            bool: True if export was successful
        """
        try:
            # Setup document styles
            self._setup_document_styles()
            
            # Add title page
            self._add_title_page()
            
            # Add table of contents
            self._add_table_of_contents()
            
            # Add executive summary
            self._add_executive_summary()
            
            # Add report sections
            self._add_report_sections()
            
            # Save the document
            self.document.save(output_path)
            
            return True
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {str(e)}")
            raise
            
    def _setup_document_styles(self):
        """Setup document styles"""
        # Title style
        title_style = self.document.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 1 style
        h1_style = self.document.styles['Heading 1']
        h1_style.font.name = 'Calibri'
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 2 style
        h2_style = self.document.styles['Heading 2']
        h2_style.font.name = 'Calibri'
        h2_style.font.size = Pt(16)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0, 51, 102)
        
        # Normal text style
        normal_style = self.document.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        
    def _add_title_page(self):
        """Add title page to document"""
        # Title
        title = self.document.add_paragraph("VMware vSphere Environment Report", style='CustomTitle')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = self.document.add_paragraph("Generated on: ")
        subtitle.add_run(self.timestamp.strftime('%Y-%m-%d %H:%M:%S'))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        self.document.add_page_break()
        
    def _add_table_of_contents(self):
        """Add table of contents to document"""
        self.document.add_heading("Table of Contents", level=1)
        
        # Add TOC placeholder
        self.document.add_paragraph("Right-click here and select 'Update Field' to update the table of contents.")
        
        # Add the TOC field
        self.document.add_paragraph().add_run("TOC \\o \"1-3\" \\h \\z \\u").font.hidden = True
        
        # Add page break
        self.document.add_page_break()
        
    def _add_executive_summary(self):
        """Add executive summary to document"""
        self.document.add_heading("Executive Summary", level=1)
        
        summary = self.document.add_paragraph()
        summary.add_run("This report contains information about the VMware vSphere environment. ")
        summary.add_run("It includes details about VMware Tools versions, VM snapshots, orphaned VMDK files, ")
        summary.add_run("and other important aspects of the environment.")
        
        # Add summary statistics
        self.document.add_heading("Summary Statistics", level=2)
        
        # Create statistics table
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Set headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Item"
        hdr_cells[1].text = "Count"
        
        # Add data rows
        items = [
            ("VMware Tools versions needing upgrade", self._count_tools_needing_upgrade()),
            ("VMs with snapshots", self._count_vms_with_snapshots()),
            ("Orphaned VMDK files", self._count_orphaned_vmdks()),
            ("Total virtual machines", self._count_total_vms()),
            ("Total ESXi hosts", self._count_total_hosts()),
            ("Total datastores", self._count_total_datastores())
        ]
        
        for item, count in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = str(count)
            
        # Add page break
        self.document.add_page_break()
        
    def _add_report_sections(self):
        """Add report sections to document"""
        # VMware Tools section (required)
        if 'vmware_tools' in self.data:
            self._add_vmware_tools_section()
            
        # Snapshots section (required)
        if 'snapshots' in self.data:
            self._add_snapshots_section()
            
        # Orphaned VMDKs section (required)
        if 'orphaned_vmdks' in self.data:
            self._add_orphaned_vmdks_section()
            
        # Additional sections
        if 'vms' in self.data:
            self._add_vms_section()
            
        if 'hosts' in self.data:
            self._add_hosts_section()
            
        if 'datastores' in self.data:
            self._add_datastores_section()
            
        if 'clusters' in self.data:
            self._add_clusters_section()
            
        if 'resource_pools' in self.data:
            self._add_resource_pools_section()
            
        if 'networks' in self.data:
            self._add_networks_section()
            
    def _add_vmware_tools_section(self):
        """Add VMware Tools section to document"""
        self.document.add_heading("VMware Tools Versions", level=1)
        
        # Add description
        self.document.add_paragraph(
            "The following table shows VMware Tools versions for all virtual machines, "
            "ordered by oldest version first."
        )
        
        # Create table
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Set headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "VM Name"
        hdr_cells[1].text = "Power State"
        hdr_cells[2].text = "Tools Status"
        hdr_cells[3].text = "Tools Version"
        
        # Add data rows
        for vm in self.data['vmware_tools']:
            row_cells = table.add_row().cells
            row_cells[0].text = vm['name']
            row_cells[1].text = vm['power_state']
            row_cells[2].text = vm['vmware_tools_status']
            row_cells[3].text = vm['vmware_tools_version']
            
        # Add recommendation paragraph
        if self._count_tools_needing_upgrade() > 0:
            self.document.add_paragraph()
            rec = self.document.add_paragraph()
            rec.add_run("Recommendation: ").bold = True
            rec.add_run(
                f"There are {self._count_tools_needing_upgrade()} virtual machines with outdated VMware Tools. "
                "It is recommended to update VMware Tools to the latest version to ensure optimal performance and compatibility."
            )
            
        # Add page break
        self.document.add_page_break()
        
    def _add_snapshots_section(self):
        """Add Snapshots section to document"""
        self.document.add_heading("VM Snapshots", level=1)
        
        # Add description
        self.document.add_paragraph(
            "The following table shows virtual machine snapshots, ordered by oldest first. "
            "Snapshots are not intended for long-term use and can impact performance if kept for extended periods."
        )
        
        if not self.data['snapshots']:
            self.document.add_paragraph("No snapshots found in the environment.")
        else:
            # Create table
            table = self.document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Set headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "VM Name"
            hdr_cells[1].text = "Snapshot Name"
            hdr_cells[2].text = "Description"
            hdr_cells[3].text = "Create Time"
            hdr_cells[4].text = "Age (Days)"
            
            # Add data rows
            for snapshot in self.data['snapshots']:
                row_cells = table.add_row().cells
                row_cells[0].text = snapshot['vm_name']
                row_cells[1].text = snapshot['name']
                row_cells[2].text = snapshot['description']
                row_cells[3].text = self._format_datetime(snapshot['create_time'])
                row_cells[4].text = str(snapshot['age_days'])
                
            # Add recommendation paragraph
            old_snapshots = [s for s in self.data['snapshots'] if s['age_days'] > 7]
            if old_snapshots:
                self.document.add_paragraph()
                rec = self.document.add_paragraph()
                rec.add_run("Recommendation: ").bold = True
                rec.add_run(
                    f"There are {len(old_snapshots)} snapshots older than 7 days. "
                    "It is recommended to consolidate or remove old snapshots to maintain optimal performance."
                )
                
        # Add page break
        self.document.add_page_break()
        
    def _add_orphaned_vmdks_section(self):
        """Add Orphaned VMDKs section to document"""
        self.document.add_heading("Orphaned VMDK Files", level=1)
        
        # Add description
        self.document.add_paragraph(
            "The following table shows VMDK files that appear to be orphaned or not associated with any registered virtual machine. "
            "These files may be consuming unnecessary storage space."
        )
        
        if not self.data['orphaned_vmdks']:
            self.document.add_paragraph("No orphaned VMDK files found in the environment.")
        else:
            # Create table
            table = self.document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Set headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Path"
            hdr_cells[1].text = "Datastore"
            hdr_cells[2].text = "Size"
            hdr_cells[3].text = "Reason"
            
            # Add data rows
            for vmdk in self.data['orphaned_vmdks']:
                row_cells = table.add_row().cells
                row_cells[0].text = vmdk['path']
                row_cells[1].text = vmdk['datastore']
                row_cells[2].text = self._format_size(vmdk['size'])
                row_cells[3].text = vmdk['reason']
                
            # Add recommendation paragraph
            if self.data['orphaned_vmdks']:
                total_size = sum(vmdk['size'] for vmdk in self.data['orphaned_vmdks'])
                self.document.add_paragraph()
                rec = self.document.add_paragraph()
                rec.add_run("Recommendation: ").bold = True
                rec.add_run(
                    f"There are {len(self.data['orphaned_vmdks'])} orphaned VMDK files consuming approximately {self._format_size(total_size)} of storage. "
                    "It is recommended to verify and remove these files to reclaim storage space."
                )
                
        # Add page break
        self.document.add_page_break()
        
    def _add_vms_section(self):
        """Add VMs section to document"""
        self.document.add_heading("Virtual Machines", level=1)
        
        # Add description
        self.document.add_paragraph(
            "The following table shows an overview of all virtual machines in the environment."
        )
        
        # Create table
        table = self.document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Set headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "VM Name"
        hdr_cells[1].text = "Power State"
        hdr_cells[2].text = "Guest OS"
        hdr_cells[3].text = "CPU"
        hdr_cells[4].text = "Memory (MB)"
        hdr_cells[5].text = "Used Space"
        
        # Add data rows
        for vm in self.data['vms']:
            row_cells = table.add_row().cells
            row_cells[0].text = vm['name']
            row_cells[1].text = vm['power_state']
            row_cells[2].text = vm['guest_full_name']
            row_cells[3].text = str(vm['num_cpu'])
            row_cells[4].text = str(vm['memory_mb'])
            row_cells[5].text = self._format_size(vm['used_space'])
            
        # Add page break
        self.document.add_page_break()
        
    def _add_hosts_section(self):
        """Add Hosts section to document"""
        self.document.add_heading("ESXi Hosts", level=1)
        
        # Add description
        self.document.add_paragraph(
            "The following table shows an overview of all ESXi hosts in the environment."
        )
        
        # Create table
        table = self.document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Set headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Host Name"
        hdr_cells[1].text = "Cluster"
        hdr_cells[2].text = "Connection State"
        hdr_cells[3].text = "CPU Model"
        hdr_cells[4].text = "CPU Cores"
        hdr_cells[5].text = "Memory (GB)"
        
        # Add data rows
        for host in self.data['hosts']:
            row_cells = table.add_row().cells
            row_cells[0].text = host['name']
            row_cells[1].text = host['cluster']
            row_cells[2].text = host['connection_state']
            row_cells[3].text = host['cpu_model']
            row_cells[4].text = str(host['cpu_cores'])
            row_cells[5].text = str(round(host['memory_size'], 2))
            
        # Add page break
        self.document.add_page_break()
        
    def _add_datastores_section(self):
        """Add Datastores section to document"""
        self.document.add_heading("Datastores", level=1)
        
        # Add description
        self.document.add_paragraph(
            "The following table shows an overview of all datastores in the environment."
        )
        
        # Create table
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Set headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Datastore Name"
        hdr_cells[1].text = "Type"
        hdr_cells[2].text = "Capacity"
        hdr_cells[3].text = "Free Space"
        hdr_cells[4].text = "Usage (%)"
        
        # Add data rows
        for datastore in self.data['datastores']:
            row_cells = table.add_row().cells
            row_cells[0].text = datastore['name']
            row_cells[1].text = datastore['type']
            row_cells[2].text = self._format_size(datastore['capacity'])
            row_cells[3].text = self._format_size(datastore['free_space'])
            row_cells[4].text = self._format_percent(datastore['usage_percent'])
            
        # Add recommendation paragraph
        high_usage_datastores = [ds for ds in self.data['datastores'] if ds['usage_percent'] > 85]
        if high_usage_datastores:
            self.document.add_paragraph()
            rec = self.document.add_paragraph()
            rec.add_run("Recommendation: ").bold = True
            rec.add_run(
                f"There are {len(high_usage_datastores)} datastores with usage above 85%. "
                "Consider adding more storage capacity or migrating VMs to balance usage."
            )
            
        # Add page break
        self.document.add_page_break()
        
    def _add_clusters_section(self):
        """Add Clusters section to document"""
        self.document.add_heading("Clusters", level=1)
        
        # Add description
        self.document.add_paragraph(
            "The following table shows an overview of all clusters in the environment."
        )
        
        # Create table
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Set headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Cluster Name"
        hdr_cells[1].text = "Hosts"
        hdr_cells[2].text = "DRS Enabled"
        hdr_cells[3].text = "HA Enabled"
        hdr_cells[4].text = "Total Memory (GB)"
        
        # Add data rows
        for cluster in self.data['clusters']:
            row_cells = table.add_row().cells
            row_cells[0].text = cluster['name']
            row_cells[1].text = str(cluster['hosts'])
            row_cells[2].text = str(cluster['drs_enabled'])
            row_cells[3].text = str(cluster['ha_enabled'])
            row_cells[4].text = str(round(cluster['total_memory'] / (1024 * 1024 * 1024), 2)) if cluster['total_memory'] else "0"
            
        # Add page break
        self.document.add_page_break()
        
    def _add_resource_pools_section(self):
        """Add Resource Pools section to document"""
        self.document.add_heading("Resource Pools", level=1)
        
        # Add description
        self.document.add_paragraph(
            "The following table shows an overview of all resource pools in the environment."
        )
        
        # Create table
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Set headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Resource Pool Name"
        hdr_cells[1].text = "Parent"
        hdr_cells[2].text = "CPU Shares"
        hdr_cells[3].text = "CPU Limit"
        hdr_cells[4].text = "Memory Limit"
        
        # Add data rows
        for pool in self.data['resource_pools']:
            row_cells = table.add_row().cells
            row_cells[0].text = pool['name']
            row_cells[1].text = f"{pool['parent_type']}: {pool['parent_name']}"
            row_cells[2].text = str(pool['cpu_shares'])
            row_cells[3].text = str(pool['cpu_limit']) if pool['cpu_limit'] != -1 else "Unlimited"
            row_cells[4].text = str(pool['memory_limit']) if pool['memory_limit'] != -1 else "Unlimited"
            
        # Add page break
        self.document.add_page_break()
        
    def _add_networks_section(self):
        """Add Networks section to document"""
        self.document.add_heading("Networks", level=1)
        
        # Add description
        self.document.add_paragraph(
            "The following table shows an overview of all networks in the environment."
        )
        
        # Create table
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Set headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Network Name"
        hdr_cells[1].text = "Type"
        hdr_cells[2].text = "Accessible"
        hdr_cells[3].text = "Additional Info"
        
        # Add data rows
        for network in self.data['networks']:
            row_cells = table.add_row().cells
            row_cells[0].text = network['name']
            row_cells[1].text = network['type']
            row_cells[2].text = str(network['accessible'])
            
            # Add additional info based on network type
            if network['type'] == 'DistributedVirtualPortgroup':
                if 'dvs_name' in network and 'vlan_id' in network:
                    row_cells[3].text = f"DVS: {network['dvs_name']}, VLAN: {network['vlan_id']}"
                else:
                    row_cells[3].text = "Distributed Virtual Portgroup"
            else:
                row_cells[3].text = "Standard Network"
                
    def _count_tools_needing_upgrade(self):
        """Count VMs with VMware Tools needing upgrade"""
        if 'vmware_tools' not in self.data:
            return 0
            
        return sum(1 for vm in self.data['vmware_tools'] 
                  if vm['vmware_tools_version'] == 'guestToolsNeedUpgrade')
                  
    def _count_vms_with_snapshots(self):
        """Count VMs with snapshots"""
        if 'snapshots' not in self.data:
            return 0
            
        # Get unique VM names
        vm_names = set(snapshot['vm_name'] for snapshot in self.data['snapshots'])
        return len(vm_names)
        
    def _count_orphaned_vmdks(self):
        """Count orphaned VMDKs"""
        if 'orphaned_vmdks' not in self.data:
            return 0
            
        return len(self.data['orphaned_vmdks'])
        
    def _count_total_vms(self):
        """Count total VMs"""
        if 'vms' in self.data:
            return len(self.data['vms'])
        return 0
        
    def _count_total_hosts(self):
        """Count total hosts"""
        if 'hosts' in self.data:
            return len(self.data['hosts'])
        return 0
        
    def _count_total_datastores(self):
        """Count total datastores"""
        if 'datastores' in self.data:
            return len(self.data['datastores'])
        return 0
        
    @staticmethod
    def _format_datetime(value):
        """Format datetime value for display"""
        if isinstance(value, datetime.datetime):
            return value.strftime('%Y-%m-%d %H:%M:%S')
        return str(value)
        
    @staticmethod
    def _format_size(value):
        """Format byte size value for display"""
        try:
            return humanize.naturalsize(value, binary=True)
        except:
            return str(value)
            
    @staticmethod
    def _format_percent(value):
        """Format percentage value for display"""
        try:
            return f"{value:.2f}%"
        except:
            return str(value)
